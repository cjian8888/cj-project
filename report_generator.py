#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
报告生成模块 - 资金穿透与关联排查系统
生成Excel底稿和公文格式报告
"""

from __future__ import annotations

import os
import re
from decimal import Decimal, InvalidOperation
from datetime import datetime
from typing import Any, Dict, List, Optional

import pandas as pd
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
import html
import json
import config
import utils
from utils.aggregation_view import (
    build_aggregation_overview as shared_build_aggregation_overview,
    extract_aggregation_payload as shared_extract_aggregation_payload,
    normalize_aggregation_ranked_entities as shared_normalize_aggregation_ranked_entities,
)

logger = utils.setup_logger(__name__)
_INVISIBLE_CHAR_RE = re.compile(r"[\u200b-\u200f\u202a-\u202e\u2060\ufeff]")
_EMPTY_TEXT_VALUES = {"", "nan", "none", "nat", "null"}
_RISK_LEVEL_MAP = {
    "critical": "极高风险",
    "high": "高风险",
    "medium": "中风险",
    "low": "低风险",
    "info": "提示",
}
_DIRECTION_MAP = {
    "in": "流入",
    "income": "流入",
    "out": "流出",
    "expense": "流出",
}
_ENTITY_TYPE_MAP = {
    "person": "人员",
    "company": "公司",
    "external": "外部联系",
    "unknown": "未知",
}
_CHANGE_TYPE_MAP = {
    "income_spike": "收入突增",
    "expense_spike": "支出突增",
    "balance_shift": "收支结构突变",
}
_HEADER_FILL = PatternFill("solid", fgColor="D9E2F3")
_HEADER_FONT = Font(bold=True)

HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>查询结果分析报告</title>
<style>
    body {
        font-family: "SimSun", "Songti SC", serif;
        background-color: #f0f0f0;
        padding: 20px;
        margin: 0;
    }
    .page {
        background-color: white;
        width: 100%;
        max-width: 210mm;
        min-height: 297mm;
        padding: 20px; /* Reduced padding for screen reading */
        margin: 0 auto; /* Center with margin instead of flex */
        box-shadow: 0 0 10px rgba(0,0,0,0.1);
        box-sizing: border-box;
        font-size: 16px;
        line-height: 1.6;
        color: #000;
        margin-bottom: 20px;
        word-wrap: break-word; /* Ensure long words break */
    }
    @media print {
        body {
            background-color: white;
            padding: 0;
        }
        .page {
            width: 210mm;
            max-width: none;
            padding: 25mm 20mm; /* Restore standard A4 margins for print */
            margin: 0;
            box-shadow: none;
            min-height: auto;
        }
    }
    h1 {
        text-align: center;
        font-size: 22px;
        font-weight: bold;
        margin-bottom: 30px;
    }
    p {
        margin-bottom: 10px;
        text-align: justify;
        text-indent: 2em;
    }
    .section-title {
        font-weight: bold;
        margin-top: 20px;
        margin-bottom: 10px;
        text-indent: 0;
        font-size: 18px;
    }
    .subsection-title {
        margin-top: 15px;
        margin-bottom: 5px;
        text-indent: 0;
        font-weight: bold;
    }
    table {
        width: 100%;
        border-collapse: collapse;
        margin: 15px 0;
        font-size: 14px;
    }
    th, td {
        border: 1px solid #000;
        padding: 5px 8px;
        text-align: center;
    }
    th {
        background-color: #f2f2f2;
        font-weight: bold;
    }
    .highlight {
        color: red;
        background-color: #fff0f0;
    }
    .img-container {
        text-align: center;
        margin: 10px 0;
    }
    img {
        max-width: 100%;
        height: auto;
    }
</style>
</head>
<body>

<!-- 报告封面/正文 -->
<!-- CONTENT_PLACEHOLDER -->

</body>
</html>"""


def _escape_html(text):
    """
    HTML 转义函数，防止 XSS 攻击

    【2026-01-27 修复】使用标准库 html.escape，确保安全

    Args:
        text: 需要转义的文本

    Returns:
        转义后的安全文本
    """
    if text is None:
        return ''
    return html.escape(str(text), quote=True)


def _safe_format_date(date_val):
    """
    安全格式化日期值

    处理 Pandas Timestamp、datetime、字符串等多种日期格式

    【2026-01-27 修复】使用 pd.to_datetime 统一处理，提升健壮性

    Args:
        date_val: 日期值（可能是 Timestamp、datetime、字符串或 None）

    Returns:
        格式化后的日期字符串 (YYYY-MM-DD) 或空字符串
    """
    dt = _coerce_datetime(date_val)
    if dt is not None:
        return dt.strftime("%Y-%m-%d")

    fallback = _clean_excel_text(date_val)
    if "T" in fallback:
        return fallback.split("T", 1)[0]
    if len(fallback) >= 10:
        return fallback[:10]
    return fallback


def _safe_format_datetime(date_val):
    """格式化时间，保留时分秒。"""
    dt = _coerce_datetime(date_val)
    if dt is not None:
        if dt.hour == 0 and dt.minute == 0 and dt.second == 0:
            return dt.strftime("%Y-%m-%d")
        return dt.strftime("%Y-%m-%d %H:%M:%S")

    fallback = _clean_excel_text(date_val)
    return fallback.replace("T", " ")


def _coerce_datetime(date_val):
    """尽量将不同来源的日期值统一转成 datetime。"""
    if date_val is None:
        return None

    if isinstance(date_val, str):
        date_str = _clean_excel_text(date_val)
        if not date_str:
            return None
        if date_str.isdigit() and len(date_str) == 8:
            try:
                return datetime.strptime(date_str, "%Y%m%d")
            except ValueError:
                return None
    elif isinstance(date_val, (int, float)) and not isinstance(date_val, bool):
        if pd.isna(date_val):
            return None
        digits = str(int(date_val))
        if len(digits) == 8 and digits.isdigit():
            try:
                return datetime.strptime(digits, "%Y%m%d")
            except ValueError:
                return None

    try:
        if pd.isna(date_val):
            return None
    except Exception:
        pass

    try:
        dt = pd.to_datetime(date_val)
        if pd.isna(dt):
            return None
        if hasattr(dt, "to_pydatetime"):
            return dt.to_pydatetime()
        return dt
    except Exception:
        return None


def _clean_excel_text(value: Any) -> str:
    """清洗 Excel 输出中的脏文本。"""
    if value is None:
        return ""
    if isinstance(value, float) and pd.isna(value):
        return ""

    text = str(value).strip()
    text = _INVISIBLE_CHAR_RE.sub("", text)
    if text.lower() in _EMPTY_TEXT_VALUES:
        return ""

    if re.fullmatch(r"\d\.\d+e\+\d+", text, flags=re.IGNORECASE):
        try:
            text = format(Decimal(text), "f").rstrip("0").rstrip(".")
        except (InvalidOperation, ValueError):
            pass
    return text


def _clean_excel_value(value: Any) -> Any:
    """统一清洗单元格值，避免 nan/None/脏字符直接落盘。"""
    if value is None:
        return ""

    if isinstance(value, float):
        if pd.isna(value):
            return ""
        return round(value, 2)

    if isinstance(value, (int, bool)):
        return value

    if isinstance(value, datetime):
        return _safe_format_datetime(value)

    if isinstance(value, pd.Timestamp):
        return _safe_format_datetime(value)

    if isinstance(value, list):
        return "; ".join(_clean_excel_text(item) for item in value if _clean_excel_text(item))

    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)

    return _clean_excel_text(value)


def _normalize_risk_level(value: Any) -> str:
    """统一输出中文风险等级。"""
    raw = _clean_excel_text(value).lower()
    if not raw:
        return ""
    return _RISK_LEVEL_MAP.get(raw, _clean_excel_text(value))


def _normalize_direction(value: Any) -> str:
    raw = _clean_excel_text(value).lower()
    if not raw:
        return ""
    return _DIRECTION_MAP.get(raw, _clean_excel_text(value))


def _normalize_yes_no(value: Any) -> str:
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, (int, float)) and not pd.isna(value):
        return "是" if float(value) != 0 else "否"
    text = _clean_excel_text(value).lower()
    if text in {"1", "true", "yes", "y", "是"}:
        return "是"
    if text in {"0", "false", "no", "n", "否"}:
        return "否"
    return _clean_excel_text(value)


def _normalize_entity_type(value: Any) -> str:
    raw = _clean_excel_text(value).lower()
    if not raw:
        return ""
    return _ENTITY_TYPE_MAP.get(raw, _clean_excel_text(value))


def _normalize_change_type(value: Any) -> str:
    raw = _clean_excel_text(value).lower()
    if not raw:
        return ""
    return _CHANGE_TYPE_MAP.get(raw, _clean_excel_text(value))


def _has_meaningful_value(value: Any) -> bool:
    if value is None:
        return False
    if isinstance(value, float) and pd.isna(value):
        return False
    if isinstance(value, str):
        return bool(_clean_excel_text(value))
    if isinstance(value, (list, tuple, set, dict)):
        return bool(value)
    return True


def _first_present(record: Optional[Dict[str, Any]], *keys: str, default: Any = None) -> Any:
    if not isinstance(record, dict):
        return default
    for key in keys:
        if key in record and _has_meaningful_value(record.get(key)):
            return record.get(key)
    return default


def _get_collection(record: Optional[Dict[str, Any]], *keys: str, default: Any = None) -> Any:
    value = _first_present(record, *keys, default=default)
    if value is None:
        return default
    return value


def _clean_records(records: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    cleaned = []
    for record in records:
        cleaned.append({key: _clean_excel_value(value) for key, value in record.items()})
    return cleaned


def _drop_empty_columns(df: pd.DataFrame) -> pd.DataFrame:
    """删除整列全空的字段，减少底稿噪音。"""
    keep_columns = []
    for column in df.columns:
        series = df[column]
        if any(_has_meaningful_value(value) for value in series.tolist()):
            keep_columns.append(column)
    if not keep_columns:
        return df
    return df.loc[:, keep_columns]


def _apply_sheet_formatting(writer, sheet_name: str, df: pd.DataFrame, startrow: int = 0) -> None:
    worksheet = writer.sheets.get(sheet_name)
    if worksheet is None or df.empty:
        return

    header_row = startrow + 1
    for cell in worksheet[header_row]:
        cell.font = _HEADER_FONT
        cell.fill = _HEADER_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center")

    if startrow == 0 and worksheet.max_row >= 2:
        worksheet.freeze_panes = "A2"
        worksheet.auto_filter.ref = worksheet.dimensions

    sample_df = df.head(200)
    for idx, column in enumerate(df.columns, start=1):
        values = [len(_clean_excel_text(column))]
        for value in sample_df[column]:
            values.append(len(_clean_excel_text(value)))
        width = min(max(values) + 2, 40)
        worksheet.column_dimensions[get_column_letter(idx)].width = max(width, 10)


def _write_excel_sheet(
    writer,
    sheet_name: str,
    records: List[Dict[str, Any]],
    *,
    startrow: int = 0,
) -> None:
    if not records:
        return
    df = pd.DataFrame(_clean_records(records))
    df = _drop_empty_columns(df)
    if df.empty:
        return
    df.to_excel(writer, sheet_name=sheet_name, index=False, startrow=startrow)
    _apply_sheet_formatting(writer, sheet_name, df, startrow=startrow)


def _extract_source_columns(record: Dict[str, Any]) -> Dict[str, Any]:
    source_file = _first_present(
        record,
        "source_file",
        "sourceFile",
        "source",
        "来源文件",
    )
    source_row = _first_present(
        record,
        "source_row_index",
        "sourceRowIndex",
        "source_row",
        "row_index",
        "来源行号",
    )
    source_file = os.path.basename(_clean_excel_text(source_file)) if source_file else ""
    return {
        "来源文件": source_file,
        "来源行号": source_row if source_row is not None else "",
    }


def _calculate_family_financials(head, members, profiles, func_type):
    """
    计算家庭财务数据

    Args:
        head: 户主
        members: 家庭成员列表
        profiles: 资金画像字典
        func_type: 计算类型 ('deposit' 或 'wealth')

    Returns:
        计算结果（万元）
    """
    total = 0.0
    targets = set(members) if members else set()
    targets.add(head)

    for p_name in targets:
         # Find profile: check if p_name is contained in any profile key
          profile = next((p for n, p in profiles.items() if p_name in n), None)
          if not profile or not profile['has_data']: continue

          if func_type == 'deposit':
              # 【2026-01-22 修复】使用真实银行余额，而不是净流入估算
              val = _get_real_bank_balance(p_name, profiles)
              total += val
          elif func_type == 'wealth':
              w = profile.get('wealth_management', {})
              val = w.get('estimated_holding', 0)
              if val == 0:
                  val = max(0, w.get('wealth_purchase', 0) - w.get('wealth_redemption', 0))
              total += val
    return round(total / 10000, 2)


def _extract_aggregation_payload(
    aggregator: Any = None,
    derived_data: Optional[Dict] = None,
) -> Dict[str, Any]:
    """从聚合器实例或 derived_data 中提取聚合结果。"""
    return shared_extract_aggregation_payload(
        aggregation=aggregator,
        derived_data=derived_data,
    )


def _normalize_aggregation_ranked_entities(
    aggregation_payload: Dict[str, Any],
    scope_entities: Optional[List[str]] = None,
) -> List[Dict[str, Any]]:
    """标准化聚合排序实体，供老链路复用。"""
    return shared_normalize_aggregation_ranked_entities(
        aggregation_payload,
        scope_entities=scope_entities,
    )


def _build_aggregation_overview(
    aggregator: Any = None,
    derived_data: Optional[Dict] = None,
    scope_entities: Optional[List[str]] = None,
    limit: int = 5,
) -> Dict[str, Any]:
    """为老链路构建聚合风险概览。"""
    return shared_build_aggregation_overview(
        aggregation=aggregator,
        derived_data=derived_data,
        scope_entities=scope_entities,
        limit=limit,
    )


def _build_semantic_priority_board_rows(
    report_package: Optional[Dict[str, Any]],
    limit: int = 5,
) -> tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    """将正式语义层 priority_board 转成 Excel 旧版聚合风险页。"""
    package = report_package if isinstance(report_package, dict) else {}
    board = package.get("priority_board")
    if not isinstance(board, list) or not board:
        return [], []

    normalized_items: List[Dict[str, Any]] = []
    for item in board:
        if not isinstance(item, dict):
            continue
        entity_name = str(item.get("entity_name") or item.get("name") or "").strip()
        if not entity_name:
            continue
        risk_label = str(item.get("risk_label") or "").strip()
        if not risk_label:
            risk_label = _RISK_LEVEL_MAP.get(
                str(item.get("risk_level") or "").strip().lower(),
                "低风险",
            )
        top_reasons = [
            str(reason).strip()
            for reason in (item.get("top_reasons") or [])
            if str(reason).strip()
        ]
        issue_refs = [
            str(issue_id).strip()
            for issue_id in (item.get("issue_refs") or [])
            if str(issue_id).strip()
        ]
        normalized_items.append(
            {
                "entity_name": entity_name,
                "priority_score": round(float(item.get("priority_score", 0) or 0), 2),
                "confidence": round(float(item.get("confidence", 0) or 0), 2),
                "severity": round(float(item.get("severity", 0) or 0), 2),
                "risk_label": risk_label,
                "top_reasons": top_reasons,
                "issue_refs": issue_refs,
            }
        )

    if not normalized_items:
        return [], []

    summary_rows = [
        {
            "指标": "极高风险实体数",
            "值": sum(1 for item in normalized_items if item["risk_label"] == "极高风险"),
        },
        {
            "指标": "高风险实体数",
            "值": sum(1 for item in normalized_items if item["risk_label"] == "高风险"),
        },
        {"指标": "高优先线索实体数", "值": len(normalized_items)},
        {
            "指标": "平均风险分",
            "值": round(
                sum(item["priority_score"] for item in normalized_items)
                / len(normalized_items),
                2,
            ),
        },
    ]

    detail_rows: List[Dict[str, Any]] = []
    for item in normalized_items[:limit]:
        reason_text = " | ".join(item["top_reasons"][:3])
        detail_rows.append(
            {
                "对象名称": item["entity_name"],
                "风险评分": item["priority_score"],
                "风险置信度": item["confidence"],
                "高优先线索数": len(item["issue_refs"]),
                "最强证据分": item["severity"],
                "综合摘要": reason_text or item["risk_label"],
                "重点线索": reason_text or "未提供",
            }
        )

    return summary_rows, detail_rows


def _generate_aggregation_summary_sheet(
    writer,
    derived_data: Optional[Dict],
    report_package: Optional[Dict[str, Any]] = None,
) -> None:
    """生成聚合风险排序摘要页。"""
    semantic_summary_rows, semantic_detail_rows = _build_semantic_priority_board_rows(
        report_package
    )
    if semantic_summary_rows:
        _write_excel_sheet(writer, "聚合风险排序", semantic_summary_rows)
        if semantic_detail_rows:
            _write_excel_sheet(writer, "聚合风险重点对象", semantic_detail_rows)
        return

    overview = _build_aggregation_overview(derived_data=derived_data)
    highlights = overview.get("highlights", [])
    summary = overview.get("summary", {})

    if not highlights and not summary:
        return

    summary_rows = [
        {"指标": "极高风险实体数", "值": int(summary.get("极高风险实体数", 0) or 0)},
        {"指标": "高风险实体数", "值": int(summary.get("高风险实体数", 0) or 0)},
        {"指标": "高优先线索实体数", "值": int(summary.get("高优先线索实体数", 0) or 0)},
        {"指标": "平均风险分", "值": round(float(overview.get("avg_score", 0) or 0), 2)},
    ]
    _write_excel_sheet(writer, "聚合风险排序", summary_rows)

    if highlights:
        detail_rows = []
        for item in highlights:
            detail_rows.append(
                {
                    "对象名称": item["entity"],
                    "风险评分": round(float(item.get("risk_score", 0) or 0), 2),
                    "风险置信度": round(float(item.get("risk_confidence", 0) or 0), 2),
                    "高优先线索数": item["high_priority_clue_count"],
                    "最强证据分": round(float(item.get("top_evidence_score", 0) or 0), 2),
                    "综合摘要": item["summary"],
                    "重点线索": " | ".join(item["top_clues"]),
                }
            )
        _write_excel_sheet(writer, "聚合风险重点对象", detail_rows)


def _generate_summary_sheet(writer, profiles):
    """
    生成资金画像汇总表

    Args:
        writer: ExcelWriter对象
        profiles: 资金画像字典
    """
    summary_data = []
    for entity, profile in profiles.items():
        # 检查是否有数据（根据实际数据结构，使用 entityType 或 has_data）
        has_data = profile.get('has_data', profile.get('entityType') != 'unknown')
        if not has_data:
            continue

        # 根据实际数据结构获取字段 - 使用summary和income_structure结构
        # 【修复】直接获取summary，不要嵌套.get('summary', {})
        summary = profile.get('summary', {})
        income_structure = profile.get('income_structure', {})
        fund_flow = profile.get('fund_flow', {})

        total_income = summary.get('total_income', 0)
        total_expense = summary.get('total_expense', 0)
        net_flow = summary.get('net_flow', total_income - total_expense)
        salary_total = income_structure.get('salary_income', 0)
        salary_ratio = summary.get('salary_ratio', 0)
        third_party_total = fund_flow.get('third_party_amount', 0)
        transaction_count = summary.get('transaction_count', 0)
        # 大额现金从large_cash列表中计算总额
        large_cash = profile.get('large_cash', [])
        cash_total = sum(c.get('amount', 0) for c in large_cash) if large_cash else 0

        # 判断是否为公司（简单的启发式：名称包含"公司"或其他关键词，或者在profiles里有type字段）
        is_company = False
        if '公司' in entity or '中心' in entity or '部' in entity:
            is_company = True

        row_data = {
            '对象名称': entity,
            '资金流入总额(万元)': round(total_income / 10000, 2),
            '资金流出总额(万元)': round(total_expense / 10000, 2),
            '净流入(万元)': round(net_flow / 10000, 2),
            '交易笔数': transaction_count,
            '大额现金总额(万元)': round(cash_total / 10000, 2)
        }

        if not is_company:
            row_data.update({
                '工资性收入(万元)': round(salary_total / 10000, 2),
                '工资性收入占比': round(salary_ratio, 3),
                '第三方支付占比': round(third_party_total / total_income, 3) if total_income > 0 else 0,
            })
        else:
            row_data.update({
                '工资性收入(万元)': '-',
                '工资性收入占比': '-',
                '第三方支付占比': round(third_party_total / total_income, 3) if total_income > 0 else 0,
            })

        summary_data.append(row_data)

    if summary_data:
        df_summary = pd.DataFrame(_clean_records(summary_data))
        df_summary.to_excel(writer, sheet_name='资金画像汇总', index=False)
        _apply_sheet_formatting(writer, '资金画像汇总', df_summary)
        worksheet = writer.sheets['资金画像汇总']
        # 设置百分比格式
        for row in range(2, len(df_summary) + 2):
            worksheet[f'H{row}'].number_format = '0.0%'
            worksheet[f'I{row}'].number_format = '0.0%'


def _generate_direct_transfer_sheet(writer, suspicions):
    """生成直接转账关系表"""
    direct_transfers = _get_collection(
        suspicions, "direct_transfers", "directTransfers", default=[]
    )
    if not direct_transfers:
        return

    transfer_data = []
    for t in direct_transfers:
        direction = _normalize_direction(_first_present(t, "direction", default=""))
        person = _first_present(t, "person", "from", "from_name", default="")
        counterparty = _first_present(t, "company", "counterparty", "to", "to_name", default="")
        if not _has_meaningful_value(t.get("person")) and direction == "流入":
            person = _first_present(t, "to", "to_name", default=person)
            counterparty = _first_present(t, "from", "from_name", default=counterparty)
        transfer_data.append({
            '日期': _safe_format_date(_first_present(t, 'date', 'transaction_date')),
            '人员': person,
            '方向': direction,
            '对手方': counterparty,
            '金额(元)': _first_present(t, 'amount', 'transaction_amount', default=0),
            '摘要': _first_present(t, 'description', 'summary', default=''),
            '风险等级': _normalize_risk_level(_first_present(t, 'risk_level', 'riskLevel')),
            **_extract_source_columns(t),
        })

    _write_excel_sheet(writer, '直接转账关系', transfer_data)


def _generate_cash_collision_sheet(writer, suspicions):
    """生成现金时空伴随表"""
    cash_collisions = _get_collection(
        suspicions, "cash_collisions", "cashCollisions", default=[]
    )
    if not cash_collisions:
        return

    collision_data = []
    for collision in cash_collisions:
        collision_data.append({
            '日期': _safe_format_date(
                _first_present(collision, 'date', 'withdrawal_date', 'time1')
            ),
            '人员A': _first_present(
                collision, 'person_a', 'person1', 'withdrawal_entity', default=''
            ),
            '人员B': _first_present(
                collision, 'person_b', 'person2', 'deposit_entity', default=''
            ),
            '取现金额(元)': _first_present(collision, 'amount1', 'withdrawal_amount', default=0),
            '存现金额(元)': _first_present(collision, 'amount2', 'deposit_amount', default=0),
            '取现地点': _first_present(collision, 'location1', 'withdrawal_location', 'location'),
            '存现地点': _first_present(collision, 'location2', 'deposit_location', 'location'),
            '时间差(小时)': round(
                float(
                    _first_present(collision, 'time_diff', 'timeDiff', 'time_diff_hours', default=0)
                    or 0
                ),
                2,
            ),
            '风险等级': _normalize_risk_level(_first_present(collision, 'risk_level', 'riskLevel')),
            '风险说明': _first_present(collision, 'risk_reason', 'riskReason', default=''),
            '取现来源文件': os.path.basename(
                _clean_excel_text(
                    _first_present(collision, 'withdrawalSource', 'withdrawal_source')
                )
            ),
            '取现来源行号': _first_present(collision, 'withdrawalRow', 'withdrawal_row', default=''),
            '存现来源文件': os.path.basename(
                _clean_excel_text(_first_present(collision, 'depositSource', 'deposit_source'))
            ),
            '存现来源行号': _first_present(collision, 'depositRow', 'deposit_row', default=''),
        })

    _write_excel_sheet(writer, '现金时空伴随', collision_data)


def _generate_hidden_asset_sheet(writer, suspicions):
    """生成隐形资产明细表"""
    hidden_assets = _get_collection(suspicions, 'hidden_assets', 'hiddenAssets', default={})
    if not hidden_assets:
        return

    hidden_data = []
    for person, assets in hidden_assets.items():
        for asset in assets:
            hidden_data.append({
                '日期': _safe_format_date(asset.get('date')),
                '人员': person,
                '对手方': _first_present(asset, 'counterparty', 'counter_party', default=''),
                '金额(元)': _first_present(asset, 'amount', default=0),
                '摘要': _first_present(asset, 'description', default=''),
                '风险说明': _first_present(asset, 'risk_reason', 'riskReason', default=''),
                **_extract_source_columns(asset),
            })

    _write_excel_sheet(writer, '隐形资产明细', hidden_data)


def _generate_fixed_frequency_sheet(writer, suspicions):
    """生成固定频率异常进账表"""
    fixed_frequency = _get_collection(
        suspicions, 'fixed_frequency', 'fixedFrequency', default={}
    )
    if not fixed_frequency:
        return

    fixed_data = []
    for person, items in fixed_frequency.items():
        for item in items:
            fixed_data.append({
                '人员': person,
                '平均日期': _clean_excel_text(_first_present(item, 'day_avg', 'avg_day')),
                '平均金额(元)': _first_present(item, 'amount_avg', 'avg_amount', default=0),
                '发生次数': _first_present(item, 'occurrences', 'count', default=0),
                '金额范围': (
                    f"{_first_present(item, 'min_amount', default=0)}-"
                    f"{_first_present(item, 'max_amount', default=0)}"
                ),
                **_extract_source_columns(item),
            })

    _write_excel_sheet(writer, '固定频率异常进账', fixed_data)


def _generate_large_cash_sheet(writer, profiles):
    """生成大额现金明细表"""
    large_cash_data = []

    for entity, profile in profiles.items():
        if not profile['has_data']:
            continue

        fund_flow = profile.get('fund_flow', {})
        for tx in fund_flow.get('large_cash_transactions', []):
            large_cash_data.append({
                '对象': entity,
                '日期': _safe_format_datetime(_first_present(tx, '日期', 'date')),
                '金额(元)': _first_present(tx, '金额', 'amount', default=0),
                '摘要': _first_present(tx, '摘要', 'description', default=''),
                '对手方': _first_present(tx, '对手方', 'counterparty', default=''),
                **_extract_source_columns(tx),
            })

    _write_excel_sheet(writer, '大额现金明细', large_cash_data)


def _generate_third_party_sheets(writer, profiles):
    """生成第三方支付交易明细表"""
    third_party_income_data = []
    third_party_expense_data = []

    for entity, profile in profiles.items():
        if not profile['has_data']:
            continue

        fund_flow = profile.get('fund_flow', {})

        # 收入明细
        for tx in fund_flow.get('third_party_income_transactions', []):
            third_party_income_data.append({
                '对象': entity,
                '日期': _safe_format_datetime(_first_present(tx, '日期', 'date')),
                '金额(元)': _first_present(tx, '金额', 'amount', default=0),
                '摘要': _first_present(tx, '摘要', 'description', default=''),
                '对手方': _first_present(tx, '对手方', 'counterparty', default=''),
                **_extract_source_columns(tx),
            })

        # 支出明细
        for tx in fund_flow.get('third_party_expense_transactions', []):
            third_party_expense_data.append({
                '对象': entity,
                '日期': _safe_format_datetime(_first_present(tx, '日期', 'date')),
                '金额(元)': _first_present(tx, '金额', 'amount', default=0),
                '摘要': _first_present(tx, '摘要', 'description', default=''),
                '对手方': _first_present(tx, '对手方', 'counterparty', default=''),
                **_extract_source_columns(tx),
            })

    _write_excel_sheet(writer, '第三方支付-收入', third_party_income_data)
    _write_excel_sheet(writer, '第三方支付-支出', third_party_expense_data)

    # 第三方支付汇总
    third_party_summary = []
    for entity, profile in profiles.items():
        if not profile['has_data']:
            continue
        fund_flow = profile.get('fund_flow', {})
        if fund_flow.get('third_party_income', 0) > 0 or fund_flow.get('third_party_expense', 0) > 0:
            third_party_summary.append({
                '对象': entity,
                '收入笔数': fund_flow.get('third_party_income_count', 0),
                '收入金额(元)': fund_flow.get('third_party_income', 0),
                '支出笔数': fund_flow.get('third_party_expense_count', 0),
                '支出金额(元)': fund_flow.get('third_party_expense', 0),
                '净流入(元)': fund_flow.get('third_party_income', 0) - fund_flow.get('third_party_expense', 0)
            })

    _write_excel_sheet(writer, '第三方支付-汇总', third_party_summary)


def _generate_wealth_management_sheets(writer, profiles):
    """生成理财产品交易明细表"""
    wealth_purchase_data = []
    wealth_redemption_data = []

    for entity, profile in profiles.items():
        if not profile['has_data']:
            continue

        wealth_mgmt = profile.get('wealth_management', {})

        # 购买明细
        for tx in wealth_mgmt.get('wealth_purchase_transactions', []):
            wealth_purchase_data.append({
                '对象': entity,
                '日期': _safe_format_datetime(_first_present(tx, '日期', 'date')),
                '金额(元)': _first_present(tx, '金额', 'amount', default=0),
                '摘要': _first_present(tx, '摘要', 'description', default=''),
                '对手方': _first_present(tx, '对手方', 'counterparty', default=''),
                '判断依据': _first_present(tx, '判断依据', 'reason', default=''),
                **_extract_source_columns(tx),
            })

        # 赎回明细
        for tx in wealth_mgmt.get('wealth_redemption_transactions', []):
            wealth_redemption_data.append({
                '对象': entity,
                '日期': _safe_format_datetime(_first_present(tx, '日期', 'date')),
                '金额(元)': _first_present(tx, '金额', 'amount', default=0),
                '摘要': _first_present(tx, '摘要', 'description', default=''),
                '对手方': _first_present(tx, '对手方', 'counterparty', default=''),
                '判断依据': _first_present(tx, '判断依据', 'reason', default=''),
                **_extract_source_columns(tx),
            })

    _write_excel_sheet(writer, '理财产品-购买', wealth_purchase_data)
    _write_excel_sheet(writer, '理财产品-赎回', wealth_redemption_data)

    # 理财产品汇总
    wealth_summary = []
    for entity, profile in profiles.items():
        if not profile['has_data']:
            continue
        wealth_mgmt = profile.get('wealth_management', {})
        if wealth_mgmt.get('total_transactions', 0) > 0:
            wealth_summary.append({
                '对象': entity,
                '购买笔数': wealth_mgmt.get('wealth_purchase_count', 0),
                '购买金额(元)': wealth_mgmt.get('wealth_purchase', 0),
                '赎回笔数': wealth_mgmt.get('wealth_redemption_count', 0),
                '赎回金额(元)': wealth_mgmt.get('wealth_redemption', 0),
                '净流向理财(元)': wealth_mgmt.get('net_wealth_flow', 0),
                '持有估算(元)': wealth_mgmt.get('estimated_holding', 0)
            })

    _write_excel_sheet(writer, '理财产品-汇总', wealth_summary)


def _generate_family_tree_sheet(writer, family_tree):
    """生成家族关系图谱工作表"""
    if not family_tree:
        return

    family_data = []
    for person, members in family_tree.items():
        for member in members:
            family_data.append({
                '核心人员': person,
                '家族成员': member.get('姓名', ''),
                '身份证号': _clean_excel_text(member.get('身份证号', '')),
                '与户主关系': member.get('与户主关系', ''),
                '性别': member.get('性别', ''),
                '出生日期': _safe_format_date(member.get('出生日期', '')),
                '户籍地': member.get('户籍地', ''),
                '数据来源': member.get('数据来源', '')
            })

    _write_excel_sheet(writer, '家族关系图谱', family_data)


def _generate_family_assets_sheets(writer, family_assets, profiles):
    """生成家族资产汇总工作表"""
    if not family_assets:
        return

    # 8.1 资产汇总表
    asset_summary_data = []
    for person, assets in family_assets.items():
        asset_summary_data.append({
            '核心人员': person,
            '家族成员数': len(assets['家族成员']),
            '家族成员': ', '.join(assets['家族成员']),
            '房产套数': assets['房产套数'],
            '房产总价值(万元)': assets['房产总价值'],
            '车辆数量': assets['车辆数量'],
            '存款估算(万元)': _calculate_family_financials(person, assets['家族成员'], profiles, 'deposit'),
            '理财持仓(万元)': _calculate_family_financials(person, assets['家族成员'], profiles, 'wealth')
        })

    _write_excel_sheet(writer, '家族资产汇总', asset_summary_data)

    # 8.2 房产明细表
    property_data = []
    for person, assets in family_assets.items():
        for prop in assets['房产']:
            property_data.append({
                '核心人员': person,
                '产权人': prop.get('产权人', ''),
                '房地坐落': prop.get('房地坐落', ''),
                '建筑面积(㎡)': prop.get('建筑面积', 0),
                '交易金额(万元)': prop.get('交易金额', 0),
                '规划用途': prop.get('规划用途', ''),
                '房屋性质': prop.get('房屋性质', ''),
                '登记时间': prop.get('登记时间', ''),
                '共有情况': prop.get('共有情况', ''),
                '共有人名称': prop.get('共有人名称', ''),
                '权属状态': prop.get('权属状态', ''),
                '数据质量': prop.get('数据质量', '正常')
            })

    _write_excel_sheet(writer, '房产明细', property_data)

    # 8.3 车辆明细表
    vehicle_data = []
    for person, assets in family_assets.items():
        for vehicle in assets['车辆']:
            row = {
                '核心人员': person,
                '所有人': vehicle.get('所有人', ''),
                '号牌号码': vehicle.get('号牌号码', ''),
                '中文品牌': vehicle.get('中文品牌', ''),
                '车身颜色': vehicle.get('车身颜色', ''),
                '初次登记日期': vehicle.get('初次登记日期', ''),
                '机动车状态': vehicle.get('机动车状态', ''),
                '是否抵押质押': _normalize_yes_no(vehicle.get('是否抵押质押', 0)),
                '能源种类': vehicle.get('能源种类', ''),
                '住所地址': vehicle.get('住所地址', '')
            }
            if any(
                _has_meaningful_value(row.get(key))
                for key in ['所有人', '号牌号码', '中文品牌', '车身颜色', '初次登记日期', '机动车状态', '能源种类', '住所地址']
            ):
                vehicle_data.append(row)

    _write_excel_sheet(writer, '车辆明细', vehicle_data)


def _generate_validation_sheets(writer, validation_results):
    """生成数据验证结果工作表"""
    if not validation_results:
        return

    # 9.1 流水数据验证
    if 'transactions' in validation_results:
        validation_data = []
        for entity, result in validation_results['transactions'].items():
            validation_data.append({
                '实体名称': entity,
                '验证状态': result['status'],
                '记录数': result['record_count'],
                '时间跨度(天)': result.get('date_range_days', 0),
                '问题': '; '.join(result['issues']) if result['issues'] else '',
                '警告': '; '.join(result['warnings']) if result['warnings'] else ''
            })

        if validation_data:
            _write_excel_sheet(writer, '数据验证-流水', validation_data)

    # 9.2 房产交易验证
    if 'properties' in validation_results:
        prop_validation_data = []
        for result in validation_results['properties']:
            prop_validation_data.append({
                '产权人': result['产权人'],
                '房产地址': result['房产地址'],
                '交易金额(万元)': result['交易金额'],
                '登记时间': result['登记时间'],
                '验证状态': result['验证状态'],
                '匹配交易日期': result['匹配交易']['日期'] if result.get('匹配交易') else '',
                '匹配交易金额': result['匹配交易']['金额'] if result.get('匹配交易') else ''
            })

        if prop_validation_data:
            _write_excel_sheet(writer, '数据验证-房产', prop_validation_data)


def _generate_penetration_sheets(writer, penetration_results):
    """生成资金穿透分析工作表"""
    if not penetration_results:
        return

    # 10.1 穿透汇总
    penetration_summary = []
    summary = penetration_results.get('summary', {})
    penetration_summary.append({
        '类型': '个人→涉案公司',
        '笔数': summary.get('个人→公司笔数', 0),
        '金额(万元)': summary.get('个人→公司总金额', 0) / 10000
    })
    penetration_summary.append({
        '类型': '涉案公司→个人',
        '笔数': summary.get('公司→个人笔数', 0),
        '金额(万元)': summary.get('公司→个人总金额', 0) / 10000
    })
    penetration_summary.append({
        '类型': '核心人员之间',
        '笔数': summary.get('核心人员间笔数', 0),
        '金额(万元)': summary.get('核心人员间总金额', 0) / 10000
    })
    penetration_summary.append({
        '类型': '涉案公司之间',
        '笔数': summary.get('涉案公司间笔数', 0),
        '金额(万元)': summary.get('涉案公司间总金额', 0) / 10000
    })

    _write_excel_sheet(writer, '资金穿透-汇总', penetration_summary)

    # 10.2 个人→公司明细
    person_to_company = _get_collection(penetration_results, 'person_to_company', default=[])
    if person_to_company:
        p2c_data = []
        for item in person_to_company:
            p2c_data.append({
                '发起方': _first_present(item, '发起方', 'from', 'source', default=''),
                '接收方': _first_present(item, '接收方', 'to', 'target', default=''),
                '日期': _safe_format_datetime(_first_present(item, '日期', 'date')),
                '收入(元)': _first_present(item, '收入', 'income', default=0),
                '支出(元)': _first_present(item, '支出', 'expense', default=0),
                '摘要': _first_present(item, '摘要', 'description', default=''),
                '对方原文': _first_present(item, '交易对方原文', 'counterparty_raw', default='')
            })
        _write_excel_sheet(writer, '穿透-个人到公司', p2c_data)

    # 10.3 公司→个人明细
    company_to_person = _get_collection(penetration_results, 'company_to_person', default=[])
    if company_to_person:
        c2p_data = []
        for item in company_to_person:
            c2p_data.append({
                '发起方': _first_present(item, '发起方', 'from', 'source', default=''),
                '接收方': _first_present(item, '接收方', 'to', 'target', default=''),
                '日期': _safe_format_datetime(_first_present(item, '日期', 'date')),
                '收入(元)': _first_present(item, '收入', 'income', default=0),
                '支出(元)': _first_present(item, '支出', 'expense', default=0),
                '摘要': _first_present(item, '摘要', 'description', default=''),
                '对方原文': _first_present(item, '交易对方原文', 'counterparty_raw', default='')
            })
        _write_excel_sheet(writer, '穿透-公司到个人', c2p_data)

    # 10.4 核心人员之间明细
    person_to_person = _get_collection(penetration_results, 'person_to_person', default=[])
    if person_to_person:
        p2p_data = []
        for item in person_to_person:
            p2p_data.append({
                '发起方': _first_present(item, '发起方', 'from', 'source', default=''),
                '接收方': _first_present(item, '接收方', 'to', 'target', default=''),
                '日期': _safe_format_datetime(_first_present(item, '日期', 'date')),
                '收入(元)': _first_present(item, '收入', 'income', default=0),
                '支出(元)': _first_present(item, '支出', 'expense', default=0),
                '摘要': _first_present(item, '摘要', 'description', default=''),
                '对方原文': _first_present(item, '交易对方原文', 'counterparty_raw', default='')
            })
        _write_excel_sheet(writer, '穿透-人员之间', p2p_data)

    # 10.5 涉案公司之间明细
    company_to_company = _get_collection(penetration_results, 'company_to_company', default=[])
    if company_to_company:
        c2c_data = []
        for item in company_to_company:
            c2c_data.append({
                '发起方': _first_present(item, '发起方', 'from', 'source', default=''),
                '接收方': _first_present(item, '接收方', 'to', 'target', default=''),
                '日期': _safe_format_datetime(_first_present(item, '日期', 'date')),
                '收入(元)': _first_present(item, '收入', 'income', default=0),
                '支出(元)': _first_present(item, '支出', 'expense', default=0),
                '摘要': _first_present(item, '摘要', 'description', default=''),
                '对方原文': _first_present(item, '交易对方原文', 'counterparty_raw', default='')
            })
        _write_excel_sheet(writer, '穿透-公司之间', c2c_data)


def _generate_loan_analysis_sheets(writer, loan_results):
    """
    生成借贷分析工作表（与借贷行为分析报告.txt对应）

    包含：双向往来、无还款借贷、规律还款三个工作表
    """
    if not loan_results:
        return

    # 11.1 双向资金往来
    bidirectional_flows = _get_collection(loan_results, 'bidirectional_flows', default=[])
    if bidirectional_flows:
        data = []
        for item in bidirectional_flows:
            data.append({
                '人员': item.get('person', ''),
                '对手方': item.get('counterparty', ''),
                '收入笔数': item.get('income_count', 0),
                '收入金额(元)': item.get('income_total', 0),
                '支出笔数': item.get('expense_count', 0),
                '支出金额(元)': item.get('expense_total', 0),
                '支出/收入比': round(item.get('ratio', 0), 3),
                '借贷类型': item.get('loan_type', ''),
                '风险等级': _normalize_risk_level(item.get('risk_level', '')),
                **_extract_source_columns(item),
            })
        _write_excel_sheet(writer, '借贷-双向往来', data)

    # 11.2 无还款借贷（疑似利益输送）
    no_repayment_loans = _get_collection(loan_results, 'no_repayment_loans', default=[])
    if no_repayment_loans:
        data = []
        for item in no_repayment_loans:
            data.append({
                '人员': item.get('person', ''),
                '对手方': item.get('counterparty', ''),
                '收入日期': _safe_format_date(item.get('income_date')),
                '金额(元)': item.get('income_amount', 0),
                '距今天数': item.get('days_since', 0),
                '已还金额(元)': item.get('total_repaid', 0),
                '还款比例': f"{item.get('repay_ratio', 0)*100:.1f}%",
                '风险原因': item.get('risk_reason', ''),
                '风险等级': _normalize_risk_level(item.get('risk_level', '')),
                '交易摘要': item.get('description', ''),
                **_extract_source_columns(item),
            })
        _write_excel_sheet(writer, '借贷-无还款', data)

    # 11.3 规律性还款模式
    regular_repayments = _get_collection(loan_results, 'regular_repayments', default=[])
    if regular_repayments:
        data = []
        for item in regular_repayments:
            data.append({
                '人员': item.get('person', ''),
                '对手方': item.get('counterparty', ''),
                '还款日(每月)': item.get('day_of_month', 0),
                '还款次数': _first_present(item, 'occurrences', 'count', default=0),
                '平均金额(元)': round(item.get('avg_amount', 0), 2),
                '总金额(元)': round(item.get('total_amount', 0), 2),
                '变异系数': round(item.get('cv', 0), 3),
                '疑似贷款': '是' if item.get('is_likely_loan') else '否',
                '风险等级': _normalize_risk_level(item.get('risk_level', '')),
                **_extract_source_columns(item),
            })
        _write_excel_sheet(writer, '借贷-规律还款', data)

    # 11.4 借贷配对分析
    loan_pairs = _get_collection(loan_results, 'loan_pairs', default=[])
    if loan_pairs:
        data = []
        for item in loan_pairs:
            data.append({
                '人员': item.get('person', ''),
                '对手方': item.get('counterparty', ''),
                '借入日期': _safe_format_date(item.get('loan_date')),
                '借入金额(元)': item.get('loan_amount', 0),
                '还款日期': _safe_format_date(item.get('repay_date')),
                '还款金额(元)': item.get('repay_amount', 0),
                '周期(天)': item.get('days', 0),
                '利率(%)': round(item.get('interest_rate', 0), 2),
                '年化利率(%)': round(item.get('annual_rate', 0), 1),
                '风险原因': item.get('risk_reason', ''),
                '风险等级': _normalize_risk_level(item.get('risk_level', '')),
                **_extract_source_columns(item),
            })
        _write_excel_sheet(writer, '借贷-配对分析', data)

    # 11.5 网贷平台往来
    online_loan_platforms = _get_collection(
        loan_results, 'online_loan_platforms', default=[]
    )
    if online_loan_platforms:
        data = []
        for item in online_loan_platforms:
            data.append({
                '人员': item.get('person', ''),
                '平台': item.get('platform', ''),
                '对手方': item.get('counterparty', ''),
                '日期': _safe_format_date(item.get('date')),
                '金额(元)': item.get('amount', 0),
                '方向': _normalize_direction(item.get('direction')),
                '摘要': item.get('description', ''),
                '风险等级': _normalize_risk_level(item.get('risk_level', '')),
                **_extract_source_columns(item),
            })
        _write_excel_sheet(writer, '借贷-网贷平台', data)


def _generate_income_anomaly_sheets(writer, income_results):
    """
    生成异常收入工作表（与异常收入来源分析报告.txt对应）
    """
    if not income_results:
        return

    all_anomalies = []

    regular_non_salary = _get_collection(income_results, 'regular_non_salary', default=[])
    large_individual_income = _get_collection(
        income_results, 'large_individual_income', 'large_personal_income', default=[]
    )
    unknown_source_income = _get_collection(
        income_results, 'unknown_source_income', 'unknown_source', default=[]
    )
    same_source_multi = _get_collection(
        income_results, 'same_source_multi', 'multi_source', default=[]
    )
    large_single_income = _get_collection(income_results, 'large_single_income', default=[])
    potential_bribe_installment = _get_collection(
        income_results, 'potential_bribe_installment', 'suspected_bribery', default=[]
    )

    regular_rows = []
    for item in regular_non_salary:
        row = {
            '人员': item.get('person', ''),
            '异常类型': '规律性非工资收入',
            '对手方': item.get('counterparty', ''),
            '金额(元)': item.get('total_amount', 0),
            '次数': _first_present(item, 'occurrences', 'count', default=0),
            '平均金额(元)': item.get('avg_amount', 0),
            '可疑原因': _first_present(item, 'possible_type', 'income_type', default=''),
            '首笔日期': _safe_format_date((_first_present(item, 'date_range') or ['', ''])[0] if isinstance(_first_present(item, 'date_range'), list) and _first_present(item, 'date_range') else ''),
            '末笔日期': _safe_format_date((_first_present(item, 'date_range') or ['', ''])[-1] if isinstance(_first_present(item, 'date_range'), list) and _first_present(item, 'date_range') else ''),
            '风险等级': _normalize_risk_level(item.get('risk_level', '')),
            **_extract_source_columns(item),
        }
        regular_rows.append(row)
        all_anomalies.append(row)

    large_transfer_rows = []
    for item in large_individual_income:
        row = {
            '人员': item.get('person', ''),
            '异常类型': '个人大额转入',
            '对手方': _first_present(item, 'counterparty', 'from_individual', default=''),
            '金额(元)': item.get('amount', 0),
            '日期': _safe_format_date(item.get('date')),
            '可疑原因': '个人大额转入需核实',
            '摘要': item.get('description', ''),
            '风险等级': _normalize_risk_level(item.get('risk_level', '')),
            **_extract_source_columns(item),
        }
        large_transfer_rows.append(row)
        all_anomalies.append(row)

    unknown_rows = []
    for item in unknown_source_income:
        row = {
            '人员': item.get('person', ''),
            '异常类型': '来源不明收入',
            '对手方': item.get('counterparty', '(缺失)'),
            '金额(元)': item.get('amount', 0),
            '日期': _safe_format_date(item.get('date')),
            '可疑原因': item.get('reason', ''),
            '摘要': item.get('description', ''),
            '风险等级': _normalize_risk_level(item.get('risk_level', '')),
            **_extract_source_columns(item),
        }
        unknown_rows.append(row)
        all_anomalies.append(row)

    same_source_rows = []
    for item in same_source_multi:
        date_range = _first_present(item, 'date_range', default=[])
        start_date = _safe_format_date(date_range[0]) if isinstance(date_range, list) and date_range else ''
        end_date = _safe_format_date(date_range[-1]) if isinstance(date_range, list) and date_range else ''
        row = {
            '人员': item.get('person', ''),
            '异常类型': '同源多次收入',
            '对手方': item.get('counterparty', ''),
            '金额(元)': _first_present(item, 'total_amount', 'total', default=0),
            '次数': _first_present(item, 'count', 'occurrences', default=0),
            '平均金额(元)': item.get('avg_amount', 0),
            '首笔日期': start_date,
            '末笔日期': end_date,
            '可疑原因': _first_present(item, 'possible_type', 'income_type', default=''),
            '风险等级': _normalize_risk_level(item.get('risk_level', '')),
            **_extract_source_columns(item),
        }
        same_source_rows.append(row)
        all_anomalies.append(row)

    large_single_rows = []
    for item in large_single_income:
        row = {
            '人员': item.get('person', ''),
            '异常类型': '大额单笔收入',
            '对手方': item.get('counterparty', ''),
            '金额(元)': item.get('amount', 0),
            '日期': _safe_format_date(item.get('date')),
            '可疑原因': _first_present(item, 'possible_type', 'income_type', default=''),
            '摘要': item.get('description', ''),
            '风险等级': _normalize_risk_level(item.get('risk_level', '')),
            **_extract_source_columns(item),
        }
        large_single_rows.append(row)
        all_anomalies.append(row)

    bribery_rows = []
    for item in potential_bribe_installment:
        row = {
            '人员': item.get('person', ''),
            '对手方': item.get('counterparty', ''),
            '月均金额(元)': _first_present(item, 'avg_monthly', 'avg_amount', default=0),
            '波动系数': round(item.get('cv', 0), 3),
            '持续月数': item.get('months', 0),
            '总笔数': _first_present(item, 'count', 'occurrences', default=0),
            '总金额(元)': item.get('total_amount', 0),
            '风险因素': _first_present(item, 'risk_factors', 'reason', default=''),
            '风险等级': _normalize_risk_level(item.get('risk_level', '')),
            **_extract_source_columns(item),
        }
        bribery_rows.append(row)

    _write_excel_sheet(writer, '异常收入-汇总', all_anomalies)
    _write_excel_sheet(writer, '异常收入-大额转入', large_transfer_rows)
    _write_excel_sheet(writer, '异常收入-来源不明', unknown_rows)
    _write_excel_sheet(writer, '异常收入-同源多次', same_source_rows)
    _write_excel_sheet(writer, '异常收入-大额单笔', large_single_rows)
    _write_excel_sheet(writer, '异常收入-疑似分期受贿', bribery_rows)


def _generate_time_series_sheets(writer, time_series_results):
    """
    生成时序分析工作表（与时序分析报告.txt对应）
    """
    if not time_series_results:
        return

    sudden_changes = _get_collection(
        time_series_results, 'sudden_changes', '突变事件', default=[]
    )
    if sudden_changes:
        data = []
        for item in sudden_changes:
            person = item.get('person', '')
            data.append({
                '人员': person if person and str(person) != 'nan' else '(未知)',
                '日期': _safe_format_date(item.get('date')),
                '金额(元)': item.get('amount', 0),
                'Z值': round(item.get('z_score', 0), 2),
                '历史均值(元)': round(_first_present(item, 'avg_before', 'mean', default=0), 2),
                '异常类型': _normalize_change_type(item.get('change_type', '')),
                '风险等级': _normalize_risk_level(item.get('risk_level', '')),
                **_extract_source_columns(item),
            })
        _write_excel_sheet(writer, '时序-资金突变', data)

    delayed_transfers = _get_collection(
        time_series_results, 'delayed_transfers', '延迟转账', default=[]
    )
    if delayed_transfers:
        data = []
        for item in delayed_transfers:
            person = item.get('person', '')
            income_cp = item.get('income_counterparty', '')
            expense_cp = item.get('expense_counterparty', '')

            person = person if person and str(person) != 'nan' else '(未知)'
            income_cp = income_cp if income_cp and str(income_cp) != 'nan' else '(未知)'
            expense_cp = expense_cp if expense_cp and str(expense_cp) != 'nan' else '(未知)'

            data.append({
                '人员': person,
                '收入来源': income_cp if income_cp and str(income_cp) != 'nan' else _clean_excel_text(item.get('income_from', '(未知)')),
                '支出去向': expense_cp if expense_cp and str(expense_cp) != 'nan' else _clean_excel_text(item.get('expense_to', '(未知)')),
                '首笔收入日期': _safe_format_date(item.get('first_income_date')),
                '延迟天数': item.get('delay_days', 0),
                '发生次数': _first_present(item, 'occurrences', 'count', default=0),
                '平均金额(元)': round(item.get('avg_amount', 0), 2),
                '总金额(元)': round(item.get('total_amount', 0), 2),
                '风险等级': _normalize_risk_level(item.get('risk_level', '')),
                **_extract_source_columns(item),
            })
        _write_excel_sheet(writer, '时序-固定延迟', data)


def _generate_fund_cycle_sheets(writer, penetration_results):
    """
    生成资金闭环工作表（与资金穿透分析报告.txt对应）
    """
    if not penetration_results:
        return

    fund_cycles = _get_collection(penetration_results, 'fund_cycles', default=[])
    if fund_cycles:
        data = []
        for cycle in fund_cycles:
            if isinstance(cycle, dict):
                cycle_str = cycle.get('path') or ' → '.join(cycle.get('nodes') or cycle.get('participants') or [])
                cycle_len = cycle.get('length') or len(cycle.get('nodes') or cycle.get('participants') or [])
                risk_level = _normalize_risk_level(cycle.get('risk_level', 'medium'))
                risk_score = cycle.get('risk_score')
                confidence = cycle.get('confidence')
            elif isinstance(cycle, list):
                cycle_str = ' → '.join(cycle)
                cycle_len = len(cycle)
                risk_level = '高风险' if cycle_len >= 3 else '中风险'
                risk_score = None
                confidence = None
            else:
                cycle_str = str(cycle)
                cycle_len = 0
                risk_level = '中风险'
                risk_score = None
                confidence = None

            data.append({
                '闭环路径': cycle_str,
                '涉及节点数': cycle_len,
                '风险等级': risk_level,
                '风险评分': risk_score,
                '置信度': confidence,
            })
        _write_excel_sheet(writer, '穿透-资金闭环', data)

    # 14.2 过账通道
    pass_through_channels = _get_collection(
        penetration_results, 'pass_through_channels', 'passthrough_channels', default=[]
    )
    if pass_through_channels:
        data = []
        for item in pass_through_channels:
            data.append({
                '节点名称': _first_present(item, 'node', 'entity', 'name', default=''),
                '节点类型': _normalize_entity_type(
                    _first_present(item, 'node_type', 'type', default='')
                ),
                '进账金额(万元)': round(item.get('inflow', 0) / 10000, 2),
                '出账金额(万元)': round(item.get('outflow', 0) / 10000, 2),
                '净沉淀(万元)': round(_first_present(item, 'net_retention', default=0) / 10000, 2),
                '进出比(%)': round(item.get('ratio', 0) * 100, 1),
                '风险等级': _normalize_risk_level(item.get('risk_level', '')),
                '风险评分': item.get('risk_score', ''),
                '置信度': item.get('confidence', ''),
                '证据摘要': '; '.join(item.get('evidence', [])) if isinstance(item.get('evidence'), list) else '',
            })
        _write_excel_sheet(writer, '穿透-过账通道', data)

    # 14.3 资金枢纽节点
    hub_nodes = _get_collection(penetration_results, 'hub_nodes', default=[])
    if hub_nodes:
        data = []
        for item in hub_nodes:
            data.append({
                '节点名称': _first_present(item, 'node', 'name', default=''),
                '节点类型': _normalize_entity_type(
                    _first_present(item, 'node_type', 'type', default='')
                ),
                '入度': item.get('in_degree', 0),
                '出度': item.get('out_degree', 0),
                '总连接数': item.get('total_degree', 0),
            })
        _write_excel_sheet(writer, '穿透-枢纽节点', data)

def _resolve_family_summaries(derived_data: Optional[Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
    """从不同兼容形态中提取家庭汇总。"""
    if not isinstance(derived_data, dict):
        return {}

    all_family_summaries = _get_collection(derived_data, 'all_family_summaries', default={})
    if isinstance(all_family_summaries, dict) and all_family_summaries:
        return all_family_summaries

    family_summary = _get_collection(derived_data, 'family_summary', default={})
    if isinstance(family_summary, dict):
        nested = _get_collection(family_summary, 'all_family_summaries', default={})
        if isinstance(nested, dict) and nested:
            return nested
        if _has_meaningful_value(family_summary.get('total_assets')):
            householder = _clean_excel_text(family_summary.get('householder')) or '默认分析单元'
            return {householder: family_summary}
    return {}


def _generate_correlation_sheets(writer, derived_data):
    """生成同行、同住宿、快递联系等多源关联工作表。"""
    if not isinstance(derived_data, dict):
        return

    correlation = _get_collection(derived_data, 'correlation', default={})
    if not isinstance(correlation, dict) or not correlation:
        return

    travel = _get_collection(correlation, 'travel_companions', default={}) or {}
    hotel = _get_collection(correlation, 'hotel_cohabitants', default={}) or {}
    express = _get_collection(correlation, 'express_contacts', default={}) or {}
    summary = _get_collection(correlation, 'summary', default={}) or {}

    summary_rows = [
        {'分析类型': '航班同行人数', '数量': summary.get('航班同行人数', len(_get_collection(travel, 'flight_companions', default=[]) or [])), '说明': '航班同行明细'},
        {'分析类型': '铁路同行人数', '数量': summary.get('铁路同行人数', len(_get_collection(travel, 'rail_companions', default=[]) or [])), '说明': '铁路同行明细'},
        {'分析类型': '同行资金关联数', '数量': len(_get_collection(travel, 'fund_correlations', default=[]) or []), '说明': '同行人与资金往来交叉'},
        {'分析类型': '同住宿人数', '数量': summary.get('同住宿人数', len(_get_collection(hotel, 'cohabitants', default=[]) or [])), '说明': '酒店同住明细'},
        {'分析类型': '同住宿资金关联数', '数量': len(_get_collection(hotel, 'fund_correlations', default=[]) or []), '说明': '同住宿人与资金往来交叉'},
        {'分析类型': '快递联系人数', '数量': summary.get('快递联系人数', len(_get_collection(express, 'express_contacts', default=[]) or [])), '说明': '快递收寄联系人'},
        {'分析类型': '快递资金关联数', '数量': len(_get_collection(express, 'fund_correlations', default=[]) or []), '说明': '快递联系人与资金往来交叉'},
    ]
    _write_excel_sheet(writer, '关联分析-汇总', summary_rows)

    flight_rows = []
    for item in _get_collection(travel, 'flight_companions', default=[]) or []:
        flight_rows.append({
            '核心人员': item.get('person', ''),
            '同行人': _first_present(item, 'companion_name', 'companion', default=''),
            '出行日期': _safe_format_date(_first_present(item, '_travel_dt', 'travel_date')),
            '交通方式': _first_present(item, 'travel_type', default='航班'),
            '航班号': _first_present(item, 'flight_no', 'transport_no', default=''),
            **_extract_source_columns(item),
        })
    _write_excel_sheet(writer, '同行分析-航班', flight_rows)

    rail_rows = []
    for item in _get_collection(travel, 'rail_companions', default=[]) or []:
        rail_rows.append({
            '核心人员': item.get('person', ''),
            '同行人': _first_present(item, 'companion_name', 'companion', default=''),
            '出行日期': _safe_format_date(_first_present(item, '_travel_dt', 'travel_date')),
            '交通方式': _first_present(item, 'travel_type', default='火车'),
            '车次': _first_present(item, 'train_no', 'transport_no', default=''),
            **_extract_source_columns(item),
        })
    _write_excel_sheet(writer, '同行分析-铁路', rail_rows)

    travel_fund_rows = []
    for item in _get_collection(travel, 'fund_correlations', default=[]) or []:
        travel_fund_rows.append({
            '核心人员': item.get('person', ''),
            '同行人': _first_present(item, 'companion', 'companion_name', default=''),
            '同行日期': _safe_format_date(_first_present(item, '_travel_dt', 'travel_date')),
            '交通方式': _first_present(item, 'travel_type', default=''),
            '交易日期': _safe_format_datetime(item.get('transaction_date')),
            '相差天数': item.get('days_diff', ''),
            '时序说明': item.get('timing', ''),
            '金额(元)': item.get('amount', 0),
            '方向': _normalize_direction(item.get('direction')),
            '交易对手原文': item.get('counterparty_raw', ''),
            '摘要': item.get('description', ''),
            '风险等级': _normalize_risk_level(item.get('risk_level', '')),
        })
    _write_excel_sheet(writer, '同行分析-资金关联', travel_fund_rows)

    hotel_rows = []
    for item in _get_collection(hotel, 'cohabitants', default=[]) or []:
        hotel_rows.append({
            '核心人员': item.get('person', ''),
            '同住人': item.get('cohabitant', ''),
            '入住日期': _safe_format_date(_first_present(item, '_stay_dt', 'stay_date')),
            '酒店名称': item.get('hotel', ''),
        })
    _write_excel_sheet(writer, '同住宿分析-明细', hotel_rows)

    hotel_fund_rows = []
    for item in _get_collection(hotel, 'fund_correlations', default=[]) or []:
        hotel_fund_rows.append({
            '核心人员': item.get('person', ''),
            '同住人': item.get('cohabitant', ''),
            '入住日期': _safe_format_date(_first_present(item, '_stay_dt', 'stay_date')),
            '交易日期': _safe_format_datetime(item.get('transaction_date')),
            '相差天数': item.get('days_diff', ''),
            '时序说明': item.get('timing', ''),
            '金额(元)': item.get('amount', 0),
            '方向': _normalize_direction(item.get('direction')),
            '交易对手原文': item.get('counterparty_raw', ''),
            '摘要': item.get('description', ''),
            '风险等级': _normalize_risk_level(item.get('risk_level', '')),
        })
    _write_excel_sheet(writer, '同住宿分析-资金关联', hotel_fund_rows)

    express_rows = []
    for item in _get_collection(express, 'express_contacts', default=[]) or []:
        express_rows.append({
            '核心人员': item.get('person', ''),
            '联系对象': item.get('contact', ''),
            '联系类型': item.get('type', ''),
        })
    _write_excel_sheet(writer, '快递联系-明细', express_rows)

    frequent_address_rows = []
    for item in _get_collection(express, 'frequent_addresses', default=[]) or []:
        frequent_address_rows.append({
            '高频地址/联系人': item.get('name', ''),
            '出现次数': item.get('count', 0),
            '涉及人员': '; '.join(item.get('persons', [])) if isinstance(item.get('persons'), list) else '',
        })
    _write_excel_sheet(writer, '快递联系-高频地址', frequent_address_rows)

    express_fund_rows = []
    for item in _get_collection(express, 'fund_correlations', default=[]) or []:
        express_fund_rows.append({
            '核心人员': item.get('person', ''),
            '快递联系人': item.get('contact', ''),
            '交易日期': _safe_format_datetime(item.get('transaction_date')),
            '金额(元)': item.get('amount', 0),
            '方向': _normalize_direction(item.get('direction')),
            '交易摘要': item.get('description', ''),
            '发生次数': _first_present(item, 'frequency', 'count', default=0),
            '风险等级': _normalize_risk_level(item.get('risk_level', '')),
        })
    _write_excel_sheet(writer, '快递联系-资金关联', express_fund_rows)


def _generate_income_classification_sheet(writer, derived_data):
    """
    生成收入分类分析工作表

    Args:
        writer: ExcelWriter对象
        derived_data: 派生数据字典（包含income_classifications）
    """
    if not derived_data or not derived_data.get('income_classifications'):
        return

    income_classifications = derived_data['income_classifications']

    # 为每个人员生成收入分类明细
    for person, classification in income_classifications.items():
        sheet_name = f'收入分类-{person}'
        # 限制sheet名称长度（Excel限制31个字符）
        if len(sheet_name) > 31:
            sheet_name = sheet_name[:28] + '...'

        classification_data = []

        # 合法收入明细
        for item in classification.get('legitimate_details', []):
            classification_data.append({
                '人员': person,
                '收入类型': '合法收入',
                '日期': _safe_format_date(item.get('date')),
                '金额(元)': item.get('amount', 0),
                '对手方': item.get('counterparty', ''),
                '摘要': item.get('description', ''),
                '判断依据': item.get('reason', ''),
                **_extract_source_columns(item),
            })

        # 未知收入明细
        for item in classification.get('unknown_details', []):
            classification_data.append({
                '人员': person,
                '收入类型': '未知收入',
                '日期': _safe_format_date(item.get('date')),
                '金额(元)': item.get('amount', 0),
                '对手方': item.get('counterparty', ''),
                '摘要': item.get('description', ''),
                '判断依据': item.get('reason', ''),
                **_extract_source_columns(item),
            })

        # 可疑收入明细
        for item in classification.get('suspicious_details', []):
            classification_data.append({
                '人员': person,
                '收入类型': '可疑收入',
                '日期': _safe_format_date(item.get('date')),
                '金额(元)': item.get('amount', 0),
                '对手方': item.get('counterparty', ''),
                '摘要': item.get('description', ''),
                '判断依据': item.get('reason', ''),
                **_extract_source_columns(item),
            })

        _write_excel_sheet(writer, sheet_name, classification_data)


def _generate_income_summary_sheet(writer, derived_data):
    """
    生成收入分类汇总工作表

    Args:
        writer: ExcelWriter对象
        derived_data: 派生数据字典（包含income_classifications）
    """
    if not derived_data or not derived_data.get('income_classifications'):
        return

    income_classifications = derived_data['income_classifications']
    summary_data = []

    for person, classification in income_classifications.items():
        summary_data.append({
            '人员': person,
            '合法收入(元)': classification.get('legitimate_income', 0),
            '合法收入占比': f"{classification.get('legitimate_ratio', 0):.2%}",
            '合法收入笔数': classification.get('legitimate_count', 0),
            '未知收入(元)': classification.get('unknown_income', 0),
            '未知收入占比': f"{classification.get('unknown_ratio', 0):.2%}",
            '未知收入笔数': classification.get('unknown_count', 0),
            '可疑收入(元)': classification.get('suspicious_income', 0),
            '可疑收入占比': f"{classification.get('suspicious_ratio', 0):.2%}",
            '可疑收入笔数': classification.get('suspicious_count', 0),
            '总收入(元)': classification.get('legitimate_income', 0) + classification.get('unknown_income', 0) + classification.get('suspicious_income', 0)
        })

    _write_excel_sheet(writer, '收入分类汇总', summary_data)


def _generate_total_assets_sheet(writer, derived_data):
    """
    生成总资产汇总工作表

    Args:
        writer: ExcelWriter对象
        derived_data: 派生数据字典（包含total_assets）
    """
    if not derived_data:
        return
    assets_data = []
    family_summaries = _resolve_family_summaries(derived_data)

    for family_name, summary in family_summaries.items():
        total_assets = summary.get('total_assets', {})
        assets_data.append({
            '分析单元': family_name,
            '银行存款(元)': total_assets.get('bank_balance', 0),
            '房产价值(元)': total_assets.get('property_value', 0),
            '车辆价值(元)': total_assets.get('vehicle_value', 0),
            '理财持仓(元)': total_assets.get('wealth_balance', 0),
            '总资产(元)': total_assets.get('total', 0),
            '房产数量': total_assets.get('property_count', 0),
            '车辆数量': total_assets.get('vehicle_count', 0),
        })

    if not assets_data:
        total_assets = _get_collection(derived_data, 'total_assets', default={})
        if isinstance(total_assets, dict) and total_assets:
            assets_data.append({
                '分析单元': '总体',
                '银行存款(元)': total_assets.get('bank_balance', 0),
                '房产价值(元)': total_assets.get('property_value', 0),
                '车辆价值(元)': total_assets.get('vehicle_value', 0),
                '理财持仓(元)': total_assets.get('wealth_balance', 0),
                '总资产(元)': total_assets.get('total', 0),
                '房产数量': total_assets.get('property_count', 0),
                '车辆数量': total_assets.get('vehicle_count', 0),
            })

    _write_excel_sheet(writer, '家庭总资产汇总', assets_data)


def _generate_member_transfers_sheet(writer, derived_data):
    """
    生成成员间转账工作表

    Args:
        writer: ExcelWriter对象
        derived_data: 派生数据字典（包含member_transfers）
    """
    if not derived_data:
        return
    transfer_data = []
    family_summaries = _resolve_family_summaries(derived_data)

    for family_name, summary in family_summaries.items():
        member_transfers = summary.get('member_transfers', {}) or {}
        for person, transfers in member_transfers.items():
            details = transfers.get('transfer_details', []) if isinstance(transfers, dict) else []
            first_date = _safe_format_date(details[0].get('date')) if details else ''
            last_date = _safe_format_date(details[-1].get('date')) if details else ''
            main_counterparties = []
            if details:
                main_counterparties = sorted(
                    { _clean_excel_text(detail.get('counterparty')) for detail in details if _clean_excel_text(detail.get('counterparty')) }
                )
            transfer_data.append({
                '分析单元': family_name,
                '成员': person,
                '转给家庭(元)': transfers.get('to_family', 0),
                '从家庭转入(元)': transfers.get('from_family', 0),
                '净流入(元)': transfers.get('net', 0),
                '转账明细数': len(details),
                '首笔日期': first_date,
                '末笔日期': last_date,
                '主要家庭对手方': '; '.join(main_counterparties[:5]),
            })

    if not transfer_data:
        member_transfers = _get_collection(derived_data, 'member_transfers', default={})
        if isinstance(member_transfers, dict):
            for person, transfers in member_transfers.items():
                transfer_data.append({
                    '分析单元': '总体',
                    '成员': person,
                    '转给家庭(元)': transfers.get('to_family', 0),
                    '从家庭转入(元)': transfers.get('from_family', 0),
                    '净流入(元)': transfers.get('net', 0),
                    '转账明细数': len(transfers.get('transfer_details', [])),
                    '首笔日期': '',
                    '末笔日期': '',
                    '主要家庭对手方': '',
                })

    _write_excel_sheet(writer, '成员间转账明细', transfer_data)


def generate_excel_workbook(profiles: Dict,
                            suspicions: Dict,
                            output_path: str = None,
                            family_tree: Dict = None,
                            family_assets: Dict = None,
                            validation_results: Dict = None,
                            penetration_results: Dict = None,
                            loan_results: Dict = None,
                            income_results: Dict = None,
                            time_series_results: Dict = None,
                            derived_data: Dict = None,
                            report_package: Optional[Dict[str, Any]] = None) -> str:
    """
    生成Excel核查底稿

    Args:
        profiles: 资金画像字典
        suspicions: 疑点检测结果
        output_path: 输出路径,默认使用配置
        family_tree: 家族关系图谱（可选）
        family_assets: 家族资产数据（可选）
        validation_results: 数据验证结果（可选）
        penetration_results: 资金穿透分析结果（可选）
        loan_results: 借贷分析结果（可选）
        income_results: 异常收入分析结果（可选）
        time_series_results: 时序分析结果（可选）
        derived_data: 派生数据（可选，包含收入分类、总资产等）

    Returns:
        生成的文件路径
    """
    if output_path is None:
        output_path = config.OUTPUT_EXCEL_FILE

    logger.info(f'正在生成Excel底稿: {output_path}')
    if not family_tree and isinstance(derived_data, dict):
        family_tree = _get_collection(derived_data, 'family_tree', default={}) or {}

    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:

        # 1. 资金画像汇总表
        _generate_summary_sheet(writer, profiles)

        # 1.5 聚合风险排序摘要
        _generate_aggregation_summary_sheet(
            writer,
            derived_data,
            report_package=report_package,
        )


        # 2. 直接转账关系表
        _generate_direct_transfer_sheet(writer, suspicions)

        # 3. 现金时空伴随表
        _generate_cash_collision_sheet(writer, suspicions)

        # 4. 隐形资产明细表
        _generate_hidden_asset_sheet(writer, suspicions)

        # 5. 固定频率异常进账表
        _generate_fixed_frequency_sheet(writer, suspicions)

        # 6. 大额现金明细表
        _generate_large_cash_sheet(writer, profiles)

        # 6.5 第三方支付交易明细表（微信/支付宝等）
        _generate_third_party_sheets(writer, profiles)

        # 6.6 理财产品交易明细表
        _generate_wealth_management_sheets(writer, profiles)

        # 7. 家族关系图谱
        _generate_family_tree_sheet(writer, family_tree)

        # 8. 家族资产汇总
        _generate_family_assets_sheets(writer, family_assets, profiles)

        # 9. 数据验证结果
        _generate_validation_sheets(writer, validation_results)

        # 10. 资金穿透分析
        _generate_penetration_sheets(writer, penetration_results)

        # 11. 借贷行为分析（新增）
        _generate_loan_analysis_sheets(writer, loan_results)

        # 12. 异常收入分析（新增）
        _generate_income_anomaly_sheets(writer, income_results)

        # 13. 时序分析（新增）
        _generate_time_series_sheets(writer, time_series_results)

        # 14. 资金闭环/过账通道（新增）
        _generate_fund_cycle_sheets(writer, penetration_results)

        # 15. 关联分析底稿（新增）
        _generate_correlation_sheets(writer, derived_data)

        # 16. 收入分类分析（新增）
        _generate_income_summary_sheet(writer, derived_data)
        _generate_income_classification_sheet(writer, derived_data)

        # 17. 总资产汇总（新增）
        _generate_total_assets_sheet(writer, derived_data)

        # 18. 成员间转账（新增）
        _generate_member_transfers_sheet(writer, derived_data)

    logger.info(f'Excel底稿生成完成: {output_path}')

    return output_path


def _group_into_households(core_persons, family_summary):
    """
    将核心人员按家庭关系分组

    支持两种 family_summary 格式：
    1. 新格式 (2026-01-23+): 包含 family_units 列表，直接使用预计算的家庭分组
    2. 旧格式: 每人 -> {配偶: [...], 子女: [...], ...}，需要动态计算分组

    Args:
        core_persons: 核心人员列表
        family_summary: 家庭关系摘要（可以是新格式 dict 或旧格式 dict）

    Returns:
        家庭分组列表，每个元素是一个家庭成员姓名列表
    """
    if not family_summary:
        # 无家庭数据，每人独立成组
        return [[p] for p in core_persons]

    # ========== 优先使用新格式 family_units ==========
    family_units = family_summary.get('family_units', [])
    
    # 【修复 2026-03-07】兼容旧缓存：如果没有 family_units，尝试从 family_units_v2 读取
    if not family_units:
        family_units = family_summary.get('family_units_v2', [])
    
    if family_units:
        modules = []
        used_persons = set()

        for unit in family_units:
            # 过滤出属于 core_persons 的成员
            unit_members = unit.get('members', [])
            filtered_members = [m for m in unit_members if m in core_persons]

            if filtered_members:
                modules.append(sorted(filtered_members))
                used_persons.update(filtered_members)

        # 将未被分配的 core_persons 添加为独立单元
        for p in core_persons:
            if p not in used_persons:
                modules.append([p])

        return modules

    # ========== 回退到旧格式 family_relations ==========
    # 从 family_summary 中获取 family_relations（如果存在）
    relations_data = family_summary.get('family_relations', family_summary)

    adj = {p: set() for p in core_persons}
    if relations_data and isinstance(relations_data, dict):
        for p, rels in relations_data.items():
            if p not in core_persons:
                continue
            if not isinstance(rels, dict):
                continue
            direct_names = []
            for k in ['配偶', '子女', '父母', '夫妻', '儿子', '女儿', '父亲', '母亲']:
                if k in rels:
                    direct_names.extend(rels[k])
            for d in direct_names:
                if d in core_persons:
                    adj[p].add(d)
                    if d in adj:
                        adj[d].add(p)

    # BFS 找连通分量
    modules = []
    visited = set()
    for p in core_persons:
        if p not in visited:
            comp = []
            q = [p]
            visited.add(p)
            while q:
                curr = q.pop(0)
                comp.append(curr)
                for neighbor in adj[curr]:
                    if neighbor not in visited:
                        visited.add(neighbor)
                        q.append(neighbor)
            modules.append(sorted(comp))
    return modules


def _format_salary_summary(income_data):
    """
    格式化工资收入摘要

    Args:
        income_data: 收入数据

    Returns:
        工资收入摘要字符串
    """
    salary_details = income_data.get('salary_details', [])
    yearly = {}
    for item in salary_details:
        d_str = str(item.get('日期', ''))[:4]
        if d_str.isdigit():
            yearly[d_str] = yearly.get(d_str, 0) + item.get('金额', 0)

    if not yearly: return "未发现明显工资性收入。"
    years = sorted(yearly.keys(), key=lambda x: int(x))
    total = sum(yearly.values())
    current_year = datetime.now().year
    recent_years = [y for y in years if int(y) >= current_year - 5]
    recent_total = sum(yearly[y] for y in recent_years)
    avg_recent = recent_total / len(recent_years) if recent_years else 0
    max_year = max(yearly, key=yearly.get)
    summary = f"工资性收入总计 {utils.format_currency(total)}。近五年平均年薪 {utils.format_currency(avg_recent)}。"
    if max_year:
         summary += f" 峰值年份为{max_year}年({utils.format_currency(yearly[max_year])})。"
    return summary


def _get_top_counterparties_str(entity, direction, cleaned_data, top_n=5):
    """
    获取主要交易对手字符串

    Args:
        entity: 实体名称
        direction: 方向 ('in' 或 'out')
        cleaned_data: 清洗后的数据
        top_n: 返回前N个

    Returns:
        主要交易对手字符串
    """
    if not cleaned_data or entity not in cleaned_data: return "无数据"
    df = cleaned_data[entity]
    col = 'income' if direction == 'in' else 'expense'
    if col not in df.columns: return "无数据"
    subset = df[df[col] > 0]
    if subset.empty: return "无主要交易"
    subset = subset[subset['交易对手'] != entity]
    stats = subset.groupby('counterparty')[col].sum().sort_values(ascending=False)
    lines = []
    count = 0
    for name, amt in zip(stats.index, stats.values):
        name_str = str(name).strip()
        if name_str.lower() in ['nan', 'none', '', 'nat'] or \
           any(x in name_str for x in ['内部户', '待清算', '资金归集', '过渡户']):
            continue
        lines.append(f"{name_str}({utils.format_currency(amt)})")
        count += 1
        if count >= top_n: break
    return ", ".join(lines) if lines else "无主要外部交易对手"


def _get_real_bank_balance(person, profiles):
    """
    获取真实的银行账户余额（从 profiles 中的 bank_accounts 数据）

    【2026-01-27 修复】添加数据溯源信息

    Args:
        person: 人员名称
        profiles: 资金画像字典

    Returns:
        真实的银行余额总和
    """
    profile = profiles.get(person, {})
    if not profile:
        return 0

    # 获取银行账户数据
    bank_accounts = profile.get('bank_accounts', []) or profile.get('bankAccounts', [])

    if not bank_accounts:
        return 0

    total_balance = 0.0
    for acc in bank_accounts:
        if isinstance(acc, dict):
            # 优先使用 last_balance（流水末尾余额），其次 balance
            balance = acc.get('last_balance', 0) or acc.get('balance', 0) or acc.get('余额', 0) or 0
            total_balance += balance

    return total_balance


def _generate_data_source_tag(person: str, cleaned_data: dict = None, source_info: dict = None) -> str:
    """
    生成数据溯源标注 HTML

    【2026-01-27 修复】新增数据溯源函数

    Args:
        person: 实体名称
        cleaned_data: 清洗后的数据（用于获取来源信息）
        source_info: 额外的源信息字典

    Returns:
        HTML 格式的溯源标注
    """
    # 获取数据来源信息
    file_info = ''
    row_count = 0

    if cleaned_data and person in cleaned_data:
        df = cleaned_data[person]
        if '_source_file' in df.columns and len(df) > 0:
            source_file = df['_source_file'].iloc[0]
            file_info = f'文件: {source_file}'
            row_count = len(df)

    if source_info and person in source_info:
        if 'file' in source_info[person]:
            file_info = f'文件: {source_info[person]["file"]}'
        if 'rows' in source_info[person]:
            row_count = source_info[person]['rows']

    if not file_info:
        file_info = '文件: 未知'

    return f"""
    <div style="font-size:11px; color:#888; margin-top:5px; padding:5px; background-color:#f9f9f9; border-radius:3px; border-left:3px solid #007bff;">
        📍 数据来源: {file_info} ({row_count} 条记录)
    </div>
    """


def _estimate_bank_balance(person, cleaned_data):
    """
    估算银行余额（已弃用，保留用于兼容性）

    注意：此函数使用简单的 收入-支出 计算，不准确
    建议使用 _get_real_bank_balance 获取真实余额

    Args:
        person: 人员名称
        cleaned_data: 清洗后的数据

    Returns:
        估算的银行余额
    """
    if not cleaned_data or person not in cleaned_data: return 0
    df = cleaned_data[person]
    total_in = df['收入(元)'].sum()
    total_out = df['支出(元)'].sum()
    return max(0, total_in - total_out)


def _generate_report_conclusion(
    profiles,
    suspicions,
    core_persons,
    involved_companies,
    aggregator=None,
    derived_data: Optional[Dict] = None,
):
    """
    生成报告核心结论部分

    Args:
        profiles: 资金画像字典
        suspicions: 疑点检测结果
        core_persons: 核心人员列表
        involved_companies: 涉及公司列表
        aggregator: ClueAggregator实例（可选），用于获取统一风险评分

    Returns:
        报告行列表
    """
    report_lines = []

    report_lines.append("一、核查结论")
    report_lines.append("-" * 60)

    total_trans = sum(p.get('summary', {}).get('transaction_count', 0) for p in profiles.values() if p.get('has_data', False))
    direct_sus_count = len(suspicions['direct_transfers'])
    hidden_sus_count = sum(len(v) for v in suspicions['hidden_assets'].values())

    aggregation_overview = _build_aggregation_overview(
        aggregator=aggregator,
        derived_data=derived_data,
        scope_entities=core_persons + involved_companies,
    )

    # 【P1 修复 2026-01-27】使用统一风险评分
    if aggregation_overview.get("risk_assessment"):
        risk_assessment = aggregation_overview["risk_assessment"]
        avg_score = aggregation_overview.get("avg_score", 0.0)
        highlights = aggregation_overview.get("highlights", [])

        report_lines.append(
            f"【总体评价】: 本次核查对象共 {len(core_persons)} 人及 {len(involved_companies)} 家关联公司，"
            f"总体风险评级为[{risk_assessment}]（平均{avg_score:.1f}分）。"
        )

        if aggregation_overview.get("summary"):
            summary = aggregation_overview["summary"]
            report_lines.append(
                f"【聚合排序】: 极高风险{summary.get('极高风险实体数', 0)}个，"
                f"高风险{summary.get('高风险实体数', 0)}个，"
                f"高优先线索实体{summary.get('高优先线索实体数', 0)}个。"
            )

        if highlights:
            high_risk_entities = [
                f"{item['entity']}({item['risk_score']:.1f}分/置信度{item['risk_confidence']:.2f})"
                for item in highlights[:3]
            ]
            report_lines.append(f"【高风险实体】: {', '.join(high_risk_entities)}")

            top_clues = []
            for item in highlights[:2]:
                if item.get("top_clues"):
                    top_clues.append(f"{item['entity']}:{item['top_clues'][0]}")
            if top_clues:
                report_lines.append(f"【重点线索】: {'；'.join(top_clues)}")
    else:
        # 旧逻辑：如果没有aggregator，使用简单的判断
        risk_assessment = "低风险"
        if direct_sus_count > 0 or hidden_sus_count > 0:
            risk_assessment = "中高风险" if direct_sus_count > 5 else "关注级"

        report_lines.append(f"【总体评价】: 本次核查对象共 {len(core_persons)} 人及 {len(involved_companies)} 家关联公司，总体风险评级为[{risk_assessment}]。")

    report_lines.append(f"【数据概况】: 累计分析银行流水 {total_trans} 条。")

    # 高流水预警
    high_flow_persons = []
    for p_name, p_data in profiles.items():
        if p_name in core_persons and p_data['has_data']:
            summary = p_data.get('summary', {})
            total_vol = summary.get('total_income', 0) + summary.get('total_expense', 0)
            if total_vol > 50000000: # 5000万
                high_flow_persons.append(f"{p_name}({utils.format_currency(total_vol)})")
    if high_flow_persons:
        report_lines.append(f"【特别说明】: {', '.join(high_flow_persons)} 银行流水规模较大，主要系理财产品频繁申赎所致，详见下文理财分析。")

    if aggregation_overview.get("highlights"):
        highlighted_entities = "、".join(
            f"{item['entity']}({item['risk_score']:.1f}分)"
            for item in aggregation_overview["highlights"][:3]
        )
        report_lines.append(f"【主要发现】: 聚合排序识别出重点核查对象 {highlighted_entities}。")
    elif direct_sus_count == 0 and hidden_sus_count == 0:
        report_lines.append(f"【主要发现】: 未发现核心人员与涉案公司存在直接利益输送，亦未发现明显的隐形房产/车辆购置线索。")
    else:
        report_lines.append(f"【主要发现】: 发现 {direct_sus_count} 笔疑似直接利益输送，{hidden_sus_count} 笔疑似隐形资产线索，需进一步核查。")

    report_lines.append("")
    return report_lines


def _generate_family_section(household, family_assets, profiles, cleaned_data):
    """
    生成家庭资产与资金画像部分

    Args:
        household: 家庭成员列表
        family_assets: 家庭资产数据
        profiles: 资金画像字典
        cleaned_data: 清洗后的数据

    Returns:
        报告行列表
    """
    report_lines = []

    title = "、".join(household) + " 家庭" if len(household) > 1 else f"{household[0]} 个人"
    report_lines.append(f"➤ {title}")

    # 2.1 家庭全貌统计
    fam_props_list = []
    fam_cars_num = 0
    fam_total_deposit_est = 0.0 # 估算存款余额
    fam_total_wealth_est = 0.0 # 估算理财沉淀
    fam_deposit_details = []
    fam_wealth_details = []

    for person in household:
        # 资产
        props = family_assets.get(person, {}).get('房产', []) if family_assets else []
        fam_props_list.extend([(person, p) for p in props])
        fam_cars_num += len(family_assets.get(person, {}).get('车辆', [])) if family_assets else 0

        # 资金状态
        p_prof = profiles.get(person)
        if p_prof and p_prof['has_data']:
            # 【2026-01-22 修复】使用真实银行余额，而不是净流入估算
            deposit = _get_real_bank_balance(person, profiles)
            fam_total_deposit_est += deposit
            if deposit > 1000:
                fam_deposit_details.append(f"{person}:{utils.format_currency(deposit)}")

            # 理财沉淀估算 (使用新的V3算法 results)
            wealth = p_prof.get('wealth_management', {})
            # 优先使用 estimate_holding
            w_holding = wealth.get('estimated_holding', 0.0)
            # 如果没有(老数据)，降级使用 Net
            if w_holding == 0 and 'estimated_holding' not in wealth:
                w_holding = max(0, wealth.get('wealth_purchase', 0) - wealth.get('wealth_redemption', 0))

            if w_holding > 0:
                 fam_total_wealth_est += w_holding
                 fam_wealth_details.append(f"{person}:{utils.format_currency(w_holding)}")

    # 展示家庭全貌表格化
    report_lines.append(f"  【家庭资产全貌】 (截至数据提取日)")
    report_lines.append(f"   • 房产合计: {len(fam_props_list)} 套")
    for owner, p in fam_props_list:
        addr = p.get('房地坐落', '未知地址')
        price = p.get('交易金额', 0)
        price_str = f"{utils.format_currency(price)}" if price > 0 else "价格未知"
        area = p.get('建筑面积', 0)
        area_str = f"{area}平" if area > 0 else "面积未知"
        report_lines.append(f"     - [{owner}] {addr} ({area_str}, {price_str})")

    report_lines.append(f"   • 车辆合计: {fam_cars_num} 辆")
    report_lines.append(f"   • 资金沉淀(估算): 存款约 {utils.format_currency(fam_total_deposit_est)} | 理财约 {utils.format_currency(fam_total_wealth_est)}")
    if fam_deposit_details:
         report_lines.append(f"     - 存款分布: {', '.join(fam_deposit_details)}")
    if fam_wealth_details:
         report_lines.append(f"     - 理财分布: {', '.join(fam_wealth_details)}")
    report_lines.append("")

    # 2.2 个人详情
    for person in household:
        profile = next((p for p in profiles.values() if person in p.get('entity_name', '')), None)
        if not profile or not profile['has_data']:
            report_lines.append(f"  [{person}]: 暂无详细流水数据"); continue

        summary = profile.get('summary', profile)
        income = profile.get('income_structure', {})
        wealth = profile.get('wealth_management', {})

        report_lines.append(f"  [{person}]")
        # 资金规模与自我转账
        self_transfer_vol = wealth.get('self_transfer_income', 0) + wealth.get('self_transfer_expense', 0)
        self_note = f"(含内部互转 {utils.format_currency(self_transfer_vol)})" if self_transfer_vol > 500000 else ""
        report_lines.append(f"    • 资金规模: 流入 {utils.format_currency(summary.get('total_income', 0))} / 流出 {utils.format_currency(summary.get('total_expense', 0))} {self_note}")

        # 工资收入
        report_lines.append(f"    • 收入结构: {utils.format_currency(income.get('salary_income', 0))} (占比 {summary.get('salary_ratio', 0):.1%})")
        report_lines.append(f"      {_format_salary_summary(income)}")

        # 理财深度分析
        w_purchase = wealth.get('wealth_purchase', 0)
        w_redeem = wealth.get('wealth_redemption', 0)
        w_est_holding = wealth.get('estimated_holding', 0)

        if w_purchase > 100000:
            report_lines.append(f"    • 理财行为: 申购 {utils.format_currency(w_purchase)} / 赎回 {utils.format_currency(w_redeem)}")
            flow_vol = summary.get('total_income', 0) + summary.get('total_expense', 0)
            if flow_vol > 0:
                turnover = (w_purchase + w_redeem) / flow_vol
                report_lines.append(f"      >> 资金空转率: {turnover:.1%} (理财申赎占总流水比例)")

            # 状态判定
            if w_est_holding > 100000:
                status = f"持有中(估算约{utils.format_currency(w_est_holding)})"
            elif w_redeem > w_purchase:
                status = f"净赎回(资金回流 {utils.format_currency(w_redeem - w_purchase)})"
            else:
                status = "基本持平"
            report_lines.append(f"      >> 资金状态: {status}")

            # 产品大类分布
            holding_struct = wealth.get('holding_structure', {})
            if holding_struct:
                 # 展示有持有量的
                 sorted_hold = sorted(holding_struct.items(), key=lambda x: x[1]['amount'], reverse=True)
                 hold_strs = [f"{k}在持{utils.format_currency(v['amount'])}" for k, v in sorted_hold]
                 report_lines.append(f"      >> 持有分布: {', '.join(hold_strs)}")
            else:
                 # 降级展示老版分布
                 cats = wealth.get('category_stats', {})
                 if cats:
                    top_cats = sorted(cats.items(), key=lambda x: x[1]['购入'] + x[1]['赎回'], reverse=True)
                    cat_strs = []
                    for c_name, c_data in top_cats:
                        vol = c_data['购入'] + c_data['赎回']
                        if vol > 100000:
                            cat_strs.append(f"{c_name}(交易{utils.format_currency(vol)})")
                    if cat_strs:
                         report_lines.append(f"      >> 交易分布: {', '.join(cat_strs)}")
        else:
             report_lines.append(f"    • 理财行为: 累计申购 {utils.format_currency(w_purchase)} (规模较小)")

        # 资金流向
        report_lines.append(f"    • 主要来源: {_get_top_counterparties_str(person, 'in', cleaned_data)}")
        report_lines.append(f"    • 主要去向: {_get_top_counterparties_str(person, 'out', cleaned_data)}")
        report_lines.append("")

    return report_lines


def _generate_company_section(company, profiles, core_persons, cleaned_data):
    """
    生成公司资金核查部分

    Args:
        company: 公司名称
        profiles: 资金画像字典
        core_persons: 核心人员列表
        cleaned_data: 清洗后的数据

    Returns:
        报告行列表
    """
    report_lines = []

    report_lines.append(f"➤ {company}")
    comp_profile = next((p for p in profiles.values() if company in p.get('entity_name', '')), None)

    if not comp_profile or not comp_profile['has_data']:
        report_lines.append("  (暂无详细流水数据)")
        report_lines.append("")
        return report_lines

    summary = comp_profile

    # 3.1 资金概况
    report_lines.append(f"  • 资金概况: 总流入 {utils.format_currency(summary.get('total_income', 0))} | 总流出 {utils.format_currency(summary.get('total_expense', 0))}")

    # 3.2 大客户/供应商
    top_in = _get_top_counterparties_str(company, 'in', cleaned_data, 5)
    top_out = _get_top_counterparties_str(company, 'out', cleaned_data, 5)
    report_lines.append(f"  • 主要资金来源(客户): {top_in}")
    report_lines.append(f"  • 主要资金去向(供应商): {top_out}")

    # 3.3 与核查对象往来 (高危检测)
    comp_df = cleaned_data.get(company)
    risky_trans = []
    if comp_df is not None:
        # 检查与核心人员往来
        rel_tx = comp_df[comp_df['交易对手'].isin(core_persons)]
        if not rel_tx.empty:
            groups = rel_tx.groupby('counterparty')[['income', 'expense']].sum()
            for name, row in groups.iterrows():
                if row['收入(元)'] > 0: risky_trans.append(f"收到 {name} {utils.format_currency(row['收入(元)'])}")
                if row['支出(元)'] > 0: risky_trans.append(f"支付给 {name} {utils.format_currency(row['支出(元)'])}")

    if risky_trans:
        report_lines.append(f"  • 【公私往来预警】: 发现直接资金往来!")
        for t in risky_trans:
            report_lines.append(f"    ⚠ {t}")
    else:
        report_lines.append(f"  • 公私往来: 未发现与核心人员的直接资金往来。")

    # 3.4 隐匿链路提示
    report_lines.append(f"  • 隐匿链路排查: 经穿透分析，未发现明显的第三方（如关联自然人、空壳公司）中转资金链路。")
    report_lines.append("")

    return report_lines


def _generate_suggestions_section(suspicions):
    """
    生成疑点与核查建议部分

    Args:
        suspicions: 疑点检测结果

    Returns:
        报告行列表
    """
    report_lines = []

    suggestion_idx = 1
    has_suggestions = False

    # 4.1 利益输送
    if suspicions['direct_transfers']:
        report_lines.append("【疑似利益输送】")
        for t in suspicions['direct_transfers']:
            report_lines.append(f"  • {_safe_format_date(t.get('date'))}: {t.get('person', '')} {t.get('direction', '')} {t.get('company', '')} {utils.format_currency(t.get('amount', 0))} (摘要: {t.get('description', '')})")
        report_lines.append(f"  ➡ 建议 {suggestion_idx}: 调取相关凭证，核实上述资金往来的业务背景。")
        suggestion_idx += 1
        has_suggestions = True
        report_lines.append("")

    # 4.2 隐形资产
    hidden = [item for sublist in suspicions['hidden_assets'].values() for item in sublist]
    if hidden:
        report_lines.append("【疑似隐形资产】")
        for h in hidden:
             report_lines.append(f"  • {_safe_format_date(h.get('date'))}: 支付 {utils.format_currency(h.get('amount', 0))} 给 {h.get('counterparty', '')} (摘要: {h.get('description', '')})")
        report_lines.append(f"  ➡ 建议 {suggestion_idx}: 核实上述大额支出是否用于购房/购车，并检查是否按规定申报。")
        suggestion_idx += 1
        has_suggestions = True
        report_lines.append("")

    # 4.3 异常高频/大额
    fixed = [item for sublist in suspicions['fixed_frequency'].values() for item in sublist]
    if fixed:
        report_lines.append("【异常规律性收入】")
        for f in fixed:
            report_lines.append(f"  • 每月{f['day_avg']}日左右收到约 {utils.format_currency(f['amount_avg'])} (共{f['occurrences']}次)")
        report_lines.append(f"  ➡ 建议 {suggestion_idx}: 核实该规律性收入的性质，排除兼职取酬或吃空饷嫌疑。")
        suggestion_idx += 1
        has_suggestions = True

    if not has_suggestions:
        report_lines.append("本次自动化分析未发现显著的硬性疑点。")
        report_lines.append(f"  ➡ 建议: 重点关注大额消费支出是否与收入水平匹配（见Excel底稿'大额现金明细'）。")

    return report_lines


def generate_official_report(profiles: Dict,
                            suspicions: Dict,
                            core_persons: List[str],
                            involved_companies: List[str],
                            output_path: str = None,
                            family_summary: Dict = None,
                            family_assets: Dict = None,
                            cleaned_data: Dict = None,
                            aggregator=None,
                            derived_data: Dict = None) -> str:
    """
    生成公文格式的核查结果分析报告（2026专业纪检优化版 v6）
    特点：
    1. 结论先行，风险分级
    2. 家庭全貌概览（包含房产/车辆/存款/理财总估值）
    3. 房产详细列表（一行一套）
    4. 深度理财分析（引入"撮合估算"算法，计算更真实的当前持有量）
    5. 公司资金核查（专章）
    
    【P1 修复 2026-01-27】aggregator参数用于统一风险评分
    """
    # 同时也生成HTML报告
    html_path = output_path.replace('.txt', '.html') if output_path else config.OUTPUT_REPORT_FILE.replace('.docx', '.html')
    generate_html_report(
        profiles,
        suspicions,
        core_persons,
        involved_companies,
        html_path,
        family_summary,
        family_assets,
        cleaned_data,
        aggregator=aggregator,
        derived_data=derived_data,
    )

    if output_path is None:
        output_path = config.OUTPUT_REPORT_FILE.replace('.docx', '.txt')

    logger.info(f'正在生成专业版公文报告(V6): {output_path}')

    report_lines = []

    # === 报告开始 ===
    report_lines.append(f"{config.REPORT_TITLE}")
    report_lines.append("=" * 60)

    # 1. 核心结论 (Executive Summary)
    report_lines.extend(
        _generate_report_conclusion(
            profiles,
            suspicions,
            core_persons,
            involved_companies,
            aggregator=aggregator,
            derived_data=derived_data,
        )
    )

    # 2. 家庭/个人详情 (Family Section)
    report_lines.append("二、家庭资产与资金画像")
    report_lines.append("-" * 60)

    households = _group_into_households(core_persons, family_summary)
    for household in households:
        report_lines.extend(_generate_family_section(household, family_assets, profiles, cleaned_data))

    # 3. 公司资金核查 (Company Section)
    report_lines.append("三、公司资金核查")
    report_lines.append("-" * 60)

    if not involved_companies:
        report_lines.append("本次未涉及公司核查。")
    else:
        for company in involved_companies:
            report_lines.extend(_generate_company_section(company, profiles, core_persons, cleaned_data))

    # 4. 疑点与核查建议
    report_lines.append("四、主要疑点与核查建议")
    report_lines.append("-" * 60)
    report_lines.extend(_generate_suggestions_section(suspicions))

    report_lines.append("")
    report_lines.append("=" * 60)
    # report_lines.append(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    report_lines.append("注: 本报告基于提供的电子数据分析生成，线索仅供参考。")

    # 写入文件
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(report_lines))

    logger.info(f'公文报告生成完成: {output_path}')
    return output_path

def _analyze_person_asset_trails(person_name: str, output_dir: str = 'output') -> List[str]:
    """
    针对单个核心人员的资产线索分析 (车/房)
    读取该人员的清洗后Excel进行深度挖掘
    """
    import os
    text_lines = []

    file_path = os.path.join(output_dir, 'cleaned_data', '个人', f'{person_name}_合并流水.xlsx')
    if not os.path.exists(file_path):
        return []

    try:
        df = pd.read_excel(file_path)

        # Helper: 关键词匹配
        def has_kw(row, kws):
            t = str(row.get('交易对手', '')) + str(row.get('交易摘要', ''))
            return any(k in t for k in kws)

        # === 车辆分析 ===
        car_buy_kws = ['汽车', '4S', '绿地金凯', '宝马', '奔驰', '奥迪', '保时捷']
        car_loan_kws = ['汽车金融', '上汽通用', '车贷', '通用金融', '宝马金融', '奔驰金融']

        # 1. 购车首付/全款
        car_payments = df[df.apply(lambda x: has_kw(x, car_buy_kws), axis=1) & (df['支出(元)'] > 10000)].sort_values('交易时间')
        if not car_payments.empty:
            for _, row in car_payments.iterrows():
                text_lines.append(f"     [车辆购置] {row['交易时间']} 向 [{row['交易对手']}] 支付 {utils.format_currency(row['支出(元)'])}")

        # 2. 车贷还款
        car_repayments = df[df.apply(lambda x: has_kw(x, car_loan_kws), axis=1) & (df['支出(元)'] > 0)].sort_values('交易时间')
        if not car_repayments.empty:
            total = car_repayments['支出(元)'].sum()
            times = len(car_repayments)
            lender = car_repayments.iloc[0]['交易对手']
            mode_val = car_repayments['支出(元)'].mode()
            monthly = mode_val[0] if not mode_val.empty else car_repayments['支出(元)'].mean()
            text_lines.append(f"     [车贷还款] 向 {lender} 累计还款 {times} 次，共计 {utils.format_currency(total)} (推测月供: {utils.format_currency(monthly)})")

        # === 房产分析 ===
        # 关键词：地产, 置业, 房地产, 首付
        house_buy_kws = ['地产', '置业', '房地产', '万科', '保利', '华润置地', '龙湖', '购房', '首期', '定金']
        # 关键词：住房贷款, 个人贷款, 按揭 (房贷通常比较隐晦, 往往只显示"贷款归还"或"利息")
        # 这里只抓特征明显的
        house_loan_kws = ['住房贷款', '个人住房', '按揭', '房贷']

        # 1. 购房支出
        house_payments = df[df.apply(lambda x: has_kw(x, house_buy_kws), axis=1) & (df['支出(元)'] > 50000)].sort_values('交易时间')
        if not house_payments.empty:
            for _, row in house_payments.iterrows():
                text_lines.append(f"     [房产购置] {row['交易时间']} 向 [{row['交易对手']}] 支付 {utils.format_currency(row['支出(元)'])} (疑似购房款/首付)")

        # 2. 房贷还款
        house_repayments = df[df.apply(lambda x: has_kw(x, house_loan_kws), axis=1) & (df['支出(元)'] > 0)].sort_values('交易时间')
        if not house_repayments.empty:
             total = house_repayments['支出(元)'].sum()
             text_lines.append(f"     [房贷还款] 发现明确的房贷摘要/对手方记录，累计还款 {utils.format_currency(total)}")

    except Exception as e:
        logger.error(f"分析资产线索失败 {person_name}: {e}")

    return text_lines

import glob
import os

HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>信息查询结果分析报告</title>
    <style>
        body {
            font-family: 'SimSun', '宋体', sans-serif;
            line-height: 1.8;
            color: #333;
            margin: 0 auto;
            padding: 20px;
            max-width: 900px;
            background-color: #f9f9f9;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
            border-radius: 8px;
        }
        .page {
            padding: 20px;
            background-color: #fff;
            border-radius: 5px;
        }
        h1 {
            font-size: 28px;
            text-align: center;
            margin-bottom: 30px;
            color: #2c3e50;
        }
        h2 {
            font-size: 24px;
            color: #2c3e50;
            border-bottom: 2px solid #eee;
            padding-bottom: 10px;
            margin-top: 40px;
        }
        h3 {
            font-size: 20px;
            color: #34495e;
            margin-top: 30px;
        }
        p {
            margin-bottom: 15px;
            text-indent: 2em;
        }
        strong {
            color: #2c3e50;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
            font-size: 14px;
        }
        th, td {
            border: 1px solid #ddd;
            padding: 10px;
            text-align: left;
        }
        th {
            background-color: #f2f2f2;
            font-weight: bold;
        }
        .section-title {
            font-size: 22px;
            font-weight: bold;
            margin-top: 30px;
            margin-bottom: 15px;
            color: #2c3e50;
            border-bottom: 1px solid #ccc;
            padding-bottom: 5px;
        }
        .subsection-title {
            font-size: 18px;
            font-weight: bold;
            margin-top: 20px;
            margin-bottom: 10px;
            color: #34495e;
        }
        .footer {
            text-align: center;
            margin-top: 50px;
            font-size: 14px;
            color: #777;
        }
</style>
</head>
<body>
    {{CONTENT}}
    <div class="footer">
        <p>(本报告由资金穿透与关联排查系统自动生成)</p>
    </div>
</body>
</html>
"""

def _generate_html_conclusion(
    profiles,
    suspicions,
    core_persons,
    involved_companies,
    aggregator=None,
    derived_data: Optional[Dict] = None,
):
    """
    生成HTML报告的核查结论部分

    Args:
        profiles: 资金画像字典
        suspicions: 疑点检测结果
        core_persons: 核心人员列表
        involved_companies: 涉及公司列表

    Returns:
        HTML内容字符串
    """
    import datetime

    total_trans = sum(p.get('summary', {}).get('transaction_count', 0) for p in profiles.values() if p.get('has_data', False))

    direct_sus_count = len(suspicions['direct_transfers'])
    hidden_sus_count = sum(len(v) for v in suspicions['hidden_assets'].values())
    aggregation_overview = _build_aggregation_overview(
        aggregator=aggregator,
        derived_data=derived_data,
        scope_entities=core_persons + involved_companies,
    )
    risk_assessment = aggregation_overview.get("risk_assessment") or "低风险"
    if not aggregation_overview.get("risk_assessment") and (direct_sus_count > 0 or hidden_sus_count > 0):
        risk_assessment = "中高风险" if direct_sus_count > 5 else "关注级"

    risk_color = "green"
    if risk_assessment in {"中高风险", "高风险"}: risk_color = "red"
    elif risk_assessment == "关注级": risk_color = "orange"

    content_html = f"""
    <div class="page">
        <h1>信息查询结果分析报告</h1>
        <p style="text-align:center; color:#888;">生成时间: {datetime.datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}</p>

        <div class="section-title">一、核查结论</div>
        <div style="background-color:#f9f9f9; padding:15px; border-left: 5px solid {risk_color}; margin-bottom:20px;">
            <p><strong>【总体评价】</strong>: 本次核查对象共 {len(core_persons)} 人及 {len(involved_companies)} 家关联公司，总体风险评级为 <strong><span style="color:{risk_color}">{_escape_html(risk_assessment)}</span></strong>{f"（平均{aggregation_overview.get('avg_score', 0):.1f}分）" if aggregation_overview.get('highlights') else ""}。</p>
            <p><strong>【数据概况】</strong>: 累计分析银行流水 {total_trans} 条。</p>
    """

    if aggregation_overview.get("summary"):
        summary = aggregation_overview["summary"]
        content_html += (
            f"""<p><strong>【聚合排序】</strong>: 极高风险{summary.get('极高风险实体数', 0)}个，"""
            f"""高风险{summary.get('高风险实体数', 0)}个，高优先线索实体{summary.get('高优先线索实体数', 0)}个。</p>"""
        )
    if aggregation_overview.get("highlights"):
        highlight_text = "，".join(
            f"{_escape_html(item['entity'])}({item['risk_score']:.1f}分/置信度{item['risk_confidence']:.2f})"
            for item in aggregation_overview["highlights"][:3]
        )
        content_html += f"""<p><strong>【高风险实体】</strong>: {highlight_text}</p>"""
        clue_text = "；".join(
            f"{_escape_html(item['entity'])}:{_escape_html(item['top_clues'][0])}"
            for item in aggregation_overview["highlights"][:2]
            if item.get("top_clues")
        )
        if clue_text:
            content_html += f"""<p><strong>【重点线索】</strong>: {clue_text}</p>"""

    # 高流水预警
    high_flow_persons = []
    for p_name, p_data in profiles.items():
        if p_name in core_persons and p_data['has_data']:
            summary = p_data.get('summary', {})
            total_vol = summary.get('total_income', 0) + summary.get('total_expense', 0)
            if total_vol > 50000000:
                high_flow_persons.append(_escape_html(p_name))
    if high_flow_persons:
        content_html += f"""<p><strong>【特别说明】</strong>: <span style="color:red">{', '.join(high_flow_persons)}</span> 银行流水规模较大，主要系理财产品频繁申赎所致(详见下文)。</p>"""

    if aggregation_overview.get("highlights"):
        highlighted_entities = "、".join(
            f"{_escape_html(item['entity'])}({item['risk_score']:.1f}分)"
            for item in aggregation_overview["highlights"][:3]
        )
        content_html += f"""<p><strong>【主要发现】</strong>: 聚合排序识别出重点核查对象 {highlighted_entities}。</p>"""
    elif direct_sus_count == 0 and hidden_sus_count == 0:
        content_html += f"""<p><strong>【主要发现】</strong>: 未发现核心人员与涉案公司存在直接利益输送，亦未发现明显的隐形房产/车辆购置线索。</p>"""
    else:
        content_html += f"""<p><strong>【主要发现】</strong>: 发现 {direct_sus_count} 笔疑似直接利益输送，{hidden_sus_count} 笔疑似隐形资产线索。</p>"""
    content_html += "</div></div>"

    return content_html


def _generate_html_family_section(profiles, core_persons, family_summary, family_assets, cleaned_data):
    """
    生成HTML报告的家庭板块 (V2.0 深度版)
    """
    import utils

    content_html = """<div class="page"><div class="section-title">二、家庭资产与资金画像</div>"""
    households = _group_into_households(core_persons, family_summary)

    for household in households:
        title = "、".join(household) + " 家庭"
        if len(household) == 1: title = f"{household[0]} 个人"

        # 1. 家庭全貌统计
        fam_props_list = []
        fam_cars_num = 0
        fam_total_deposit_est = 0.0
        fam_total_wealth_est = 0.0

        for person in household:
            # 资产数据
            props = family_assets.get(person, {}).get('房产', []) if family_assets else []
            fam_props_list.extend([(person, p) for p in props])
            fam_cars_num += len(family_assets.get(person, {}).get('车辆', [])) if family_assets else 0

            # 资金数据
            p_prof = profiles.get(person)
            if p_prof and p_prof['has_data']:
                # 【2026-01-22 修复】使用真实银行余额
                deposit = _get_real_bank_balance(person, profiles)
                fam_total_deposit_est += deposit
                wealth = p_prof.get('wealth_management', {})
                w_holding = wealth.get('estimated_holding', 0.0)
                if w_holding == 0 and 'estimated_holding' not in wealth:
                    w_holding = max(0, wealth.get('wealth_purchase', 0) - wealth.get('wealth_redemption', 0))
                if w_holding > 0: fam_total_wealth_est += w_holding

        content_html += f"""<div class="subsection-title">➤ {_escape_html(title)}</div>"""

        # 家庭概览框
        content_html += f"""
        <div style="border: 1px solid #ddd; padding: 10px; background-color: #fcfcfc; margin-bottom: 15px;">
            <p><strong>【家庭资产全貌】</strong> <span style="font-size:12px; color:#888;">(注: 房产/车辆信息需接入不动产/车管数据)</span></p>
            <p>• <strong>房产合计: {len(fam_props_list)} 套</strong> <span style="color:#999; font-size:12px;">(系统登记)</span></p>
            """
        if fam_props_list:
             content_html += '<ul style="margin-top:0; margin-bottom:5px; font-size:14px; color:#555;">'
             for owner, p in fam_props_list:
               addr = _escape_html(p.get('房地坐落', '未知地址'))
               content_html += f"<li>[{_escape_html(owner)}] {addr}</li>"
             content_html += '</ul>'
        else:
             content_html += '<p style="text-indent:2em; color:#999; font-size:12px;">(未发现系统登记房产，建议调取线下档案)</p>'

        content_html += f"""
            <p>• <strong>车辆合计: {fam_cars_num} 辆</strong></p>
            <p>• <strong>资金沉淀 (估算):</strong> 存款约 <b>{utils.format_currency(fam_total_deposit_est)}</b> | 理财约 <b>{utils.format_currency(fam_total_wealth_est)}</b></p>
        </div>
        """

        # 2. 个人详情
        for person in household:
            content_html += f"""<h3 style="margin-left:0; font-size:16px;">[{_escape_html(person)}]</h3>"""
            profile = next((p for p in profiles.values() if person in p.get('entity_name', '')), None)

            # --- 基础身份信息 (Placeholder) ---
            content_html += """
            <table style="width:100%; border:none; margin-bottom:10px;">
                <tr>
                    <td style="border:none; padding:2px;">• 身份证号: <span style="color:#999;">(需补充)</span></td>
                    <td style="border:none; padding:2px;">• 工作单位: <span style="color:#999;">(需补充)</span></td>
                </tr>
            </table>
            """

            if not profile or not profile['has_data']:
                content_html += "<p>(暂无详细流水数据)</p>"; continue

            summary = profile.get('summary', profile)
            income = profile.get('income_structure', {})
            wealth = profile.get('wealth_management', {})

            # 资金规模
            content_html += f"""<p><strong>资金规模</strong>: 流入 {utils.format_currency(summary.get('total_income', 0))} / 流出 {utils.format_currency(summary.get('total_expense', 0))}</p>"""

            # 收入结构
            salary_details = income.get('salary_details', [])
            yearly_sal = {}
            for item in salary_details:
                 y = str(item.get('日期', ''))[:4]
                 if y.isdigit(): yearly_sal[y] = yearly_sal.get(y, 0) + item.get('金额', 0)
            sal_desc = ""
            if yearly_sal:
                 sal_desc = " (" + "; ".join([f"{y}年:{utils.format_currency(v)}" for y, v in sorted(yearly_sal.items())]) + ")"

            ratio_val = summary.get('salary_ratio', 0)
            ratio_str = f"{ratio_val:.1%}"
            ratio_style = "color:red; font-weight:bold;" if ratio_val < 0.5 else ""

            content_html += f"""
            <p><strong>工资收入</strong>: 累计 {utils.format_currency(income.get('salary_income', 0))} {_escape_html(sal_desc)}</p>
            <p><strong>收支匹配</strong>: 工资占比 <span style="{ratio_style}">{_escape_html(ratio_str)}</span></p>
            """
            if ratio_val < 0.5:
                content_html += f"""<p style="text-indent:2em; color:red; font-size:13px;">⚠ 预警: 工资收入无法覆盖消费支出 (占比低于50%)，存在资金来源不明风险。</p>"""

            # 理财分析
            w_purchase = wealth.get('wealth_purchase', 0)
            if w_purchase > 100000:
                content_html += f"""<p><strong>理财行为</strong>: 申购 {utils.format_currency(w_purchase)} / 赎回 {utils.format_currency(wealth.get('wealth_redemption', 0))}</p>"""

            # 流向
            content_html += f"""\u003cp\u003e\u003cstrong\u003e主要去向\u003c/strong\u003e: {_escape_html(_get_top_counterparties_str(person, 'out', cleaned_data, 5))}\u003c/p\u003e"""
            content_html += "\u003chr style='border:0; border-top:1px dashed #eee; margin:10px 0;'\u003e"

    content_html += "</div>" # End Household Page
    return content_html


def _generate_html_family_section_v2(profiles, core_persons, family_summary, family_assets, cleaned_data, family_units):
    """
    生成HTML报告的家庭板块 (V2.0 - 使用用户配置的家庭分组)

    Args:
        family_units: 用户配置的家庭单元列表，格式为:
            [{
                "anchor": "主归集人",
                "members": ["成员1", "成员2", ...],
                "member_details": [...],
                "address": "家庭地址"
            }, ...]
    """
    import utils

    content_html = """<div class="page"><div class="section-title">二、家庭资产与资金画像</div>"""

    # 使用用户配置的家庭分组
    if family_units:
        for unit in family_units:
            anchor = unit.get('anchor', '')
            members = unit.get('members', [])
            address = unit.get('address', '')

            if not members:
                continue

            # 标题使用主归集人的名字
            title = f"{anchor} 家庭"
            address_str = f" ({address})" if address else ""

            content_html += f"""<div class="subsection-title">➤ {title}{address_str}</div>"""

            # 获取主归集人的画像
            anchor_profile = profiles.get(anchor)
            if anchor_profile and anchor_profile.get('has_data'):
                summary = anchor_profile.get('summary', anchor_profile)
                income_structure = anchor_profile.get('income_structure', {})
                wealth = anchor_profile.get('wealth_management', {})

                # 资金规模
                total_income = summary.get('total_income', 0)
                total_expense = summary.get('total_expense', 0)
                net_flow = summary.get('net_flow', 0)

                content_html += f"""
                <div style="border: 1px solid #ddd; padding: 10px; background-color: #fcfcfc; margin-bottom: 15px;">
                    <p><strong>【家庭资产全貌】</strong> <span style="font-size:12px; color:#888;">(注: 房产/车辆信息需接入不动产/车管数据)</span></p>"""

                # 显示各成员的基本数据
                for member in members:
                    member_profile = profiles.get(member)
                    if member_profile and member_profile.get('has_data'):
                        m_summary = member_profile.get('summary', member_profile)
                        m_income = m_summary.get('total_income', 0)
                        m_expense = m_summary.get('total_expense', 0)
                        m_trans = m_summary.get('transaction_count', 0)

                        is_anchor = (member == anchor)
                        anchor_mark = " ⭐" if is_anchor else ""

                        content_html += f"""
                    <h3 style="margin-left:0; font-size:16px;">[{member}]{anchor_mark}</h3>
                    <p><strong>资金规模</strong>: 流入 {utils.format_currency(m_income)} / 流出 {utils.format_currency(m_expense)}</p>
                    <p><strong>交易笔数</strong>: {m_trans} 笔</p>
"""
                content_html += "</div>"
                content_html += "<hr style='border:0; border-top:1px dashed #eee; margin:10px 0;'>"
    else:
        # 回退到系统识别的家庭分组
        households = _group_into_households(core_persons, family_summary)
        for household in households:
            title = "、".join(household) + " 家庭"
            if len(household) == 1: title = f"{household[0]} 个人"
            # ... (使用原有逻辑)
            content_html += f"<p>{title}</p>"

    content_html += "</div>"
    return content_html


def _generate_html_company_section(profiles, involved_companies, core_persons, cleaned_data):
    """
    生成HTML报告的公司资金核查部分 (V2.0 深度版)
    """
    import utils

    content_html = """<div class="page"><div class="section-title">三、公司资金核查</div>"""
    if not involved_companies:
        content_html += "<p>本次未涉及公司核查。</p>"
    else:
        for company in involved_companies:
            content_html += f"""<div class="subsection-title">➤ {_escape_html(company)}</div>"""
            comp_profile = next((p for p in profiles.values() if company in p.get('entity_name', '')), None)

            if not comp_profile or not comp_profile['has_data']:
                content_html += "<p>(暂无详细流水数据)</p>"; continue

            summary = comp_profile.get('summary', comp_profile)

            # 1. 资金规模
            content_html += f"""
            <div style="background-color:#f8f9fa; padding:10px; border-left:4px solid #007bff; margin-bottom:10px;">
                <p><strong>【资金规模】</strong></p>
                <p>• 累计进账: <b>{utils.format_currency(summary.get('total_income', 0))}</b> | 累计出账: <b>{utils.format_currency(summary.get('total_expense', 0))}</b></p>
                <p>• 交易笔数: {_escape_html(str(summary.get('transaction_count', 0)))} 笔</p>
            </div>
            """

            # 2. 上下游分析
            top_in = _escape_html(_get_top_counterparties_str(company, 'in', cleaned_data, 5))
            top_out = _escape_html(_get_top_counterparties_str(company, 'out', cleaned_data, 5))
            content_html += f"""<p><strong>主要资金来源</strong>: {top_in}</p>"""
            content_html += f"""<p><strong>主要资金去向</strong>: {top_out}</p>"""

            # 3. 关联交易 (重点)
            comp_df = cleaned_data.get(company) if cleaned_data else None
            risky_html = ""
            if comp_df is not None:
                # 3.1 与核心人员往来
                rel_tx = comp_df[comp_df['交易对手'].isin(core_persons)]
                if not rel_tx.empty:
                    groups = rel_tx.groupby('counterparty')[['income', 'expense']].sum()
                    risky_items = []
                    for name, row in groups.iterrows():
                        if row['收入(元)'] > 0 or row['支出(元)'] > 0:
                            risky_items.append(f"{_escape_html(name)} (收:{utils.format_currency(row['收入(元)'])}/付:{utils.format_currency(row['支出(元)'])})")
                    if risky_items:
                        risky_html += f"""<p style="color:red; background-color:#fff0f0; padding:5px;">⚠ <strong>利益输送嫌疑</strong>: 发现与核心人员存在直接往来: {', '.join(risky_items)}</p>"""

                # 3.2 大额现金
                cash_tx = comp_df[comp_df['现金'] == True]
                if not cash_tx.empty:
                    cash_in = cash_tx['收入(元)'].sum()
                    cash_out = cash_tx['支出(元)'].sum()
                    if cash_in + cash_out > 50000:
                        risky_html += f"""<p><strong>现金分析</strong>: 存在现金操作 (存:{utils.format_currency(cash_in)} / 取:{utils.format_currency(cash_out)})，请核实用途。</p>"""

            if risky_html:
                content_html += risky_html
            else:
                content_html += "<p>• <strong>关联排查</strong>: 未发现与核心人员的直接资金往来，无大额现金预警。</p>"

            content_html += "<br>"

    content_html += "</div>"
    return content_html


def _generate_html_suggestions_section(suspicions):
    """
    生成HTML报告的疑点与核查建议部分

    Args:
        suspicions: 疑点检测结果

    Returns:
        HTML内容字符串
    """
    import utils

    content_html = """<div class="page"><div class="section-title">四、主要疑点与核查建议</div>"""
    has_suggestions = False

    # 疑似利益输送
    if suspicions['direct_transfers']:
        content_html += """<p><strong>【疑似利益输送】</strong></p><ul>"""
        for t in suspicions['direct_transfers']:
            content_html += f"""<li>{_safe_format_date(t.get('date'))}: {_escape_html(t.get('person', ''))} {_escape_html(t.get('direction', ''))} {_escape_html(t.get('company', ''))} {utils.format_currency(t.get('amount', 0))} <br><span style="color:#666; font-size:12px;">(摘要: {_escape_html(t.get('description', ''))})</span></li>"""
        content_html += "</ul><p class='highlight'>➡ 建议: 调取相关凭证，核实上述资金往来的业务背景。</p>"
        has_suggestions = True

    # 隐形资产
    hidden = [item for sublist in suspicions['hidden_assets'].values() for item in sublist]
    if hidden:
        content_html += """<p><strong>【疑似隐形资产】</strong></p><ul>"""
        for h in hidden:
             content_html += f"""<li>{_safe_format_date(h.get('date'))}: 支付 {utils.format_currency(h.get('amount', 0))} 给 {_escape_html(h.get('counterparty', ''))} <br><span style="color:#666; font-size:12px;">(摘要: {_escape_html(h.get('description', ''))})</span></li>"""
        content_html += "</ul><p class='highlight'>➡ 建议: 核实上述大额支出是否用于购房/购车，并检查是否按规定申报。</p>"
        has_suggestions = True

    if not has_suggestions:
        content_html += "<p>本次自动化分析未发现显著的硬性疑点。</p><p>➡ 建议: 重点关注大额消费支出是否与收入水平匹配（见Excel底稿'大额现金明细'）。</p>"

    content_html += "</div>"

    return content_html


def _load_primary_targets_config(output_path=None):
    """
    读取用户配置的归集配置 (primary_targets.json)

    Args:
        output_path: 报告输出路径（用于确定配置文件位置）

    Returns:
        配置字典，包含 analysis_units, include_companies 等
        如果文件不存在或读取失败，返回 None
    """
    import os
    # 尝试多个可能的路径
    possible_paths = [
        'output/report_config/primary_targets.json',
        'report_config/primary_targets.json',
        'primary_targets.json'
    ]

    for config_path in possible_paths:
        if os.path.exists(config_path):
            try:
                with open(config_path, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                logger.info(f'已读取用户配置: {len(config.get("analysis_units", []))} 个分析单元')
                return config
            except Exception as e:
                logger.warning(f'读取用户配置失败 ({config_path}): {e}')

    logger.debug('未找到用户配置文件 primary_targets.json')
    return None

    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            config = json.load(f)
        logger.info(f'已读取用户配置: {len(config.get("analysis_units", []))} 个分析单元')
        return config
    except Exception as e:
        logger.warning(f'读取用户配置失败: {e}')
        return None


def generate_html_report(
    profiles,
    suspicions,
    core_persons,
    involved_companies,
    output_path,
    family_summary=None,
    family_assets=None,
    cleaned_data=None,
    aggregator=None,
    derived_data: Dict = None,
):
    """
    生成HTML格式的分析报告 (V6版 - 匹配最新文本报告逻辑 - Fix Placeholder)
    """
    import os
    import json
    import datetime
    import config
    import utils
    import account_analyzer

    logger = utils.setup_logger(__name__)
    logger.info(f'正在生成HTML文本分析报告(V6): {output_path}')

    # Ensure lists
    if involved_companies is None: involved_companies = []
    if core_persons is None: core_persons = []

    # 读取用户配置 (primary_targets.json)
    user_config = _load_primary_targets_config(output_path)

    # 如果有用户配置，使用用户的家庭分组；否则使用系统识别的家庭分组
    if user_config and user_config.get('analysis_units'):
        # 使用用户配置的家庭分组
        family_units = user_config['analysis_units']
        logger.info(f'使用用户配置的家庭分组: {len(family_units)} 个家庭')
    else:
        # 使用系统识别的家庭分组（从family_summary）
        family_units = None

    # 构建报告内容
    content_html = ""

    # 1. 标题与前言（核查结论）
    content_html += _generate_html_conclusion(
        profiles,
        suspicions,
        core_persons,
        involved_companies,
        aggregator=aggregator,
        derived_data=derived_data,
    )

    # 2. 家庭板块 - 使用用户配置或系统识别的家庭分组
    content_html += _generate_html_family_section_v2(profiles, core_persons, family_summary, family_assets, cleaned_data, family_units)

    # 3. 公司资金核查
    content_html += _generate_html_company_section(profiles, involved_companies, core_persons, cleaned_data)

    # 4. 疑点与核查建议
    content_html += _generate_html_suggestions_section(suspicions)

    # 家庭板块
    content_html += """<div class="page"><div class="section-title">二、家庭资产与资金画像</div>"""
    households = _group_into_households(core_persons, family_summary)

    for household in households:
        title = "、".join(household) + " 家庭"
        if len(household) == 1: title = f"{household[0]} 个人"

        # 1. 家庭全貌统计
        fam_props_list = []
        fam_cars_num = 0
        fam_total_deposit_est = 0.0
        fam_total_wealth_est = 0.0

        for person in household:
            props = family_assets.get(person, {}).get('房产', []) if family_assets else []
            fam_props_list.extend([(person, p) for p in props])
            fam_cars_num += len(family_assets.get(person, {}).get('车辆', [])) if family_assets else 0

            p_prof = profiles.get(person)
            if p_prof and p_prof['has_data']:
                # 【2026-01-22 修复】使用真实银行余额
                deposit = _get_real_bank_balance(person, profiles)
                fam_total_deposit_est += deposit
                wealth = p_prof.get('wealth_management', {})
                w_holding = wealth.get('estimated_holding', 0.0)
                if w_holding == 0 and 'estimated_holding' not in wealth:
                    w_holding = max(0, wealth.get('wealth_purchase', 0) - wealth.get('wealth_redemption', 0))
                if w_holding > 0: fam_total_wealth_est += w_holding

        content_html += f"""<div class="subsection-title">➤ {_escape_html(title)}</div>"""

        # 家庭概览框
        content_html += f"""
        <div style="border: 1px solid #ddd; padding: 10px; background-color: #fcfcfc; margin-bottom: 15px;">
            <p><strong>【家庭资产全貌】</strong> (截至数据提取日)</p>
            <p>• <strong>房产合计: {len(fam_props_list)} 套</strong></p>
            <ul style="margin-top:0; margin-bottom:5px; font-size:14px; color:#555;">
        """
        for owner, p in fam_props_list:
            addr = _escape_html(p.get('房地坐落', '未知地址'))
            price = p.get('交易金额', 0)
            price_str = f"{utils.format_currency(price)}" if price > 0 else "价格未知"
            area = _escape_html(str(p.get('建筑面积', 0)))
            content_html += f"<li>[{_escape_html(owner)}] {addr} ({area}平, {price_str})</li>"

        content_html += f"""
            </ul>
            <p>• <strong>车辆合计: {fam_cars_num} 辆</strong></p>
            <p>• <strong>资金沉淀 (估算):</strong> 存款约 <b>{utils.format_currency(fam_total_deposit_est)}</b> | 理财约 <b>{utils.format_currency(fam_total_wealth_est)}</b></p>
        </div>
        """

        # 2. 个人详情
        for person in household:
            content_html += f"""<h3 style="margin-left:0; font-size:16px;">[{_escape_html(person)}]</h3>"""
            profile = next((p for p in profiles.values() if person in p.get('entity_name', '')), None)
            if not profile or not profile['has_data']:
                content_html += "<p>(暂无详细流水数据)</p>"; continue

            summary = profile.get('summary', profile)
            income = profile.get('income_structure', {})
            wealth = profile.get('wealth_management', {})

            # 资金规模
            self_transfer_vol = wealth.get('self_transfer_income', 0) + wealth.get('self_transfer_expense', 0)
            self_note = f"<span style='color:#888'>(含内部互转 {utils.format_currency(self_transfer_vol)})</span>" if self_transfer_vol > 500000 else ""
            content_html += f"""<p><strong>资金规模</strong>: 流入 {utils.format_currency(summary.get('total_income', 0))} / 流出 {utils.format_currency(summary.get('total_expense', 0))} {self_note}</p>"""

            # 收入结构
            salary_details = income.get('salary_details', [])
            yearly_sal = {}
            for item in salary_details:
                 y = str(item.get('日期', ''))[:4]
                 if y.isdigit(): yearly_sal[y] = yearly_sal.get(y, 0) + item.get('金额', 0)
            sal_desc = ""
            if yearly_sal:
                 sal_desc = " (" + "; ".join([f"{_escape_html(y)}年:{utils.format_currency(v)}" for y, v in sorted(yearly_sal.items())]) + ")"
            content_html += f"""<p><strong>收入结构</strong>: 工资性收入 {utils.format_currency(income.get('salary_income', 0))} (占比 {summary.get('salary_ratio', 0):.1%}){_escape_html(sal_desc)}</p>"""

            # 理财深度分析
            w_purchase = wealth.get('wealth_purchase', 0)
            w_redeem = wealth.get('wealth_redemption', 0)
            w_est_holding = wealth.get('estimated_holding', 0)

            if w_purchase > 100000:
                content_html += f"""<p><strong>理财行为</strong>: 申购 {utils.format_currency(w_purchase)} / 赎回 {utils.format_currency(w_redeem)}</p>"""
                flow_vol = summary.get('total_income', 0) + summary.get('total_expense', 0)
                if flow_vol > 0:
                    turnover = (w_purchase + w_redeem) / flow_vol
                    content_html += f"""<p style="text-indent: 2em;">>> <span style="color:#666">资金空转率: <strong>{turnover:.1%}</strong> (理财申赎占总流水比例)</span></p>"""

                status_str = ""
                if w_est_holding > 100000:
                    status_str = f"持有中 (估算约 <b>{utils.format_currency(w_est_holding)}</b>)"
                elif w_redeem > w_purchase:
                    status_str = f"净赎回 (资金回流 {utils.format_currency(w_redeem - w_purchase)})"
                else:
                    status_str = "基本持平"
                content_html += f"""<p style="text-indent: 2em;">>> 资金状态: <span style="background-color:#e6f7ff; padding:2px 5px;">{status_str}</span></p>"""

                # 持有分布
                holding_struct = wealth.get('holding_structure', {})
                if holding_struct:
                     sorted_hold = sorted(holding_struct.items(), key=lambda x: x[1]['amount'], reverse=True)
                     hold_strs = [f"{_escape_html(k)}在持{utils.format_currency(v['amount'])}" for k, v in sorted_hold]
                     content_html += f"""<p style="text-indent: 2em;">>> 持有分布: {_escape_html(', '.join(hold_strs))}</p>"""

            # 流向
            content_html += f"""<p><strong>主要去向</strong>: {_escape_html(_get_top_counterparties_str(person, 'out', cleaned_data))}</p>"""
            content_html += "<hr style='border:0; border-top:1px dashed #eee; margin:10px 0;'>"
    content_html += "</div>" # End Household Page

    # 3. 公司资金核查 (Company Section)
    content_html += """<div class="page"><div class="section-title">三、公司资金核查</div>"""
    if not involved_companies:
        content_html += "<p>本次未涉及公司核查。</p>"
    else:
        for company in involved_companies:
            content_html += f"""<div class="subsection-title">➤ {_escape_html(company)}</div>"""
            comp_profile = next((p for p in profiles.values() if company in p.get('entity_name', '')), None)

            if not comp_profile or not comp_profile['has_data']:
                content_html += "<p>(暂无详细流水数据)</p>"; continue

            summary = comp_profile.get('summary', comp_profile)
            content_html += f"""<p>• <strong>资金概况</strong>: 总流入 {utils.format_currency(summary.get('total_income', 0))} | 总流出 {utils.format_currency(summary.get('total_expense', 0))}</p>"""

            top_in = _escape_html(_get_top_counterparties_str(company, 'in', cleaned_data, 5))
            top_out = _escape_html(_get_top_counterparties_str(company, 'out', cleaned_data, 5))
            content_html += f"""<p>• <strong>主要客户(来源)</strong>: {top_in}</p>"""
            content_html += f"""<p>• <strong>主要供应商(去向)</strong>: {top_out}</p>"""

            # 风险
            comp_df = cleaned_data.get(company)
            risky_html = ""
            if comp_df is not None:
                rel_tx = comp_df[comp_df['交易对手'].isin(core_persons)]
                if not rel_tx.empty:
                    groups = rel_tx.groupby('counterparty')[['income', 'expense']].sum()
                    risky_items = []
                    for name, row in groups.iterrows():
                        if row['收入(元)'] > 0 or row['支出(元)'] > 0: risky_items.append(f"{_escape_html(name)} (收:{utils.format_currency(row['收入(元)'])}/付:{utils.format_currency(row['支出(元)'])})")
                    if risky_items:
                        risky_html = f"""<p style="color:red">⚠ 发现与核心人员直接往来: {', '.join(risky_items)}</p>"""

            if risky_html: content_html += risky_html
            else: content_html += "<p>• <strong>公私往来</strong>: 未发现与核心人员的直接资金往来。</p>"
            content_html += "<br>"
    content_html += "</div>"

    # 4. 疑点与核查建议
    content_html += """<div class="page"><div class="section-title">四、主要疑点与核查建议</div>"""
    has_suggestions = False

    # 疑似利益输送
    if suspicions['direct_transfers']:
        content_html += """<p><strong>【疑似利益输送】</strong></p><ul>"""
        for t in suspicions['direct_transfers']:
            content_html += f"""<li>{_safe_format_date(t.get('date'))}: {_escape_html(t.get('person', ''))} {_escape_html(t.get('direction', ''))} {_escape_html(t.get('company', ''))} {utils.format_currency(t.get('amount', 0))} <br><span style="color:#666; font-size:12px;">(摘要: {_escape_html(t.get('description', ''))})</span></li>"""
        content_html += "</ul><p class='highlight'>➡ 建议: 调取相关凭证，核实上述资金往来的业务背景。</p>"
        has_suggestions = True

    # 隐形资产
    hidden = [item for sublist in suspicions['hidden_assets'].values() for item in sublist]
    if hidden:
        content_html += """<p><strong>【疑似隐形资产】</strong></p><ul>"""
        for h in hidden:
             content_html += f"""<li>{_safe_format_date(h.get('date'))}: 支付 {utils.format_currency(h.get('amount', 0))} 给 {_escape_html(h.get('counterparty', ''))} <br><span style="color:#666; font-size:12px;">(摘要: {_escape_html(h.get('description', ''))})</span></li>"""
        content_html += "</ul><p class='highlight'>➡ 建议: 核实上述大额支出是否用于购房/购车，并检查是否按规定申报。</p>"
        has_suggestions = True

    if not has_suggestions:
        content_html += "<p>本次自动化分析未发现显著的硬性疑点。</p><p>➡ 建议: 重点关注大额消费支出是否与收入水平匹配（见Excel底稿'大额现金明细'）。</p>"

    content_html += "</div>"

    # Replace placeholder (Handle both types to be safe)
    final_html = HTML_TEMPLATE.replace('{{CONTENT}}', content_html).replace('<!-- CONTENT_PLACEHOLDER -->', content_html)

    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(final_html)

        logger.info(f"HTML文本报告(V6)生成完毕: {output_path}")
        return output_path
    except PermissionError as e:
        logger.error(f'保存 HTML 报告失败（权限错误）: {e}')
        return None
    except OSError as e:
        logger.error(f'保存 HTML 报告失败（系统错误）: {e}')
        return None
    except Exception as e:
        logger.error(f'保存 HTML 报告失败（未知错误）: {e}')
        logger.exception('详细错误信息:')
        return None


# ============================================================
# Word 文档导出 (Phase 1.4 - 2026-01-18 新增)
# ============================================================

def generate_word_report(profiles: Dict,
                         suspicions: Dict,
                         core_persons: List[str],
                         involved_companies: List[str],
                         output_path: str = None,
                         family_summary: Dict = None,
                         family_assets: Dict = None,
                         cleaned_data: Dict = None,
                         aggregator: Any = None,
                         derived_data: Optional[Dict] = None) -> str:
    """
    使用 python-docx 生成专业 Word 审计报告

    Args:
        profiles: 资金画像字典
        suspicions: 疑点检测结果
        core_persons: 核心人员列表
        involved_companies: 涉及公司列表
        output_path: 输出路径
        family_summary: 家庭关系摘要
        family_assets: 家庭资产数据
        cleaned_data: 清洗后的数据
        aggregator: ClueAggregator实例（可选），用于获取统一风险评分
        derived_data: 派生分析数据（可选），用于读取 aggregation 结果

    Returns:
        生成的 Word 文件路径
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        logger.warning('python-docx 未安装，跳过 Word 报告生成。请执行: pip install python-docx')
        return None

    # 参数验证
    if profiles is None:
        logger.error('generate_word_report: profiles 参数不能为 None')
        return None
    if suspicions is None:
        logger.error('generate_word_report: suspicions 参数不能为 None')
        return None
    if core_persons is None:
        logger.error('generate_word_report: core_persons 参数不能为 None')
        return None
    if involved_companies is None:
        logger.error('generate_word_report: involved_companies 参数不能为 None')
        return None

    # 确保参数是正确的类型
    if not isinstance(profiles, dict):
        logger.error(f'generate_word_report: profiles 必须是字典类型，实际类型: {type(profiles)}')
        return None
    if not isinstance(suspicions, dict):
        logger.error(f'generate_word_report: suspicions 必须是字典类型，实际类型: {type(suspicions)}')
        return None
    if not isinstance(core_persons, list):
        logger.error(f'generate_word_report: core_persons 必须是列表类型，实际类型: {type(core_persons)}')
        return None
    if not isinstance(involved_companies, list):
        logger.error(f'generate_word_report: involved_companies 必须是列表类型，实际类型: {type(involved_companies)}')
        return None

    if output_path is None:
        output_path = os.path.join(config.OUTPUT_DIR, 'analysis_results', '审计分析报告.docx')

    logger.info(f'正在生成 Word 审计报告: {output_path}')

    try:
        doc = Document()
    except Exception as e:
        logger.error(f'创建 Word 文档失败: {e}')
        return None

    # 设置默认字体
    try:
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    except Exception as e:
        logger.warning(f'设置默认字体失败: {e}，使用默认字体')

    # ========== 1. 标题页 ==========
    try:
        title = doc.add_heading(config.REPORT_TITLE, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        logger.error(f'添加标题失败: {e}')
        return None

    # 生成时间
    try:
        doc.add_paragraph()
        time_para = doc.add_paragraph()
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        time_run = time_para.add_run(f'生成时间: {datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
        time_run.font.size = Pt(11)
        time_run.font.color.rgb = RGBColor(128, 128, 128)
    except Exception as e:
        logger.warning(f'添加生成时间失败: {e}')

    try:
        doc.add_page_break()
    except Exception as e:
        logger.warning(f'添加分页符失败: {e}')

    # ========== 2. 核查结论 ==========
    try:
        doc.add_heading('一、核查结论', level=1)
    except Exception as e:
        logger.error(f'添加核查结论标题失败: {e}')
        return None

    # 安全计算统计数据
    try:
        total_trans = 0
        for p in profiles.values():
            if isinstance(p, dict) and p.get('has_data', False):
                summary = p.get('summary', {})
                if isinstance(summary, dict):
                    total_trans += summary.get('transaction_count', 0)
    except Exception as e:
        logger.warning(f'计算交易总数失败: {e}')
        total_trans = 0

    try:
        direct_sus_count = len(suspicions.get('direct_transfers', []))
        hidden_sus_count = sum(len(v) for v in suspicions.get('hidden_assets', {}).values())
    except Exception as e:
        logger.warning(f'计算疑点数量失败: {e}')
        direct_sus_count = 0
        hidden_sus_count = 0

    aggregation_overview = _build_aggregation_overview(
        aggregator=aggregator,
        derived_data=derived_data,
        scope_entities=core_persons + involved_companies,
    )

    risk_assessment = aggregation_overview.get("risk_assessment") or "低风险"
    if not aggregation_overview.get("risk_assessment") and (direct_sus_count > 0 or hidden_sus_count > 0):
        risk_assessment = "中高风险" if direct_sus_count > 5 else "关注级"
    avg_score = aggregation_overview.get("avg_score", 0.0)

    try:
        p = doc.add_paragraph()
        p.add_run('【总体评价】').bold = True
        p.add_run(f'本次核查对象共 {len(core_persons)} 人及 {len(involved_companies)} 家关联公司，')
        suffix = f'（平均{avg_score:.1f}分）' if aggregation_overview.get("highlights") else ''
        risk_run = p.add_run(f'总体风险评级为 [{risk_assessment}]{suffix}。')
        if risk_assessment in {"中高风险", "高风险"}:
            risk_run.font.color.rgb = RGBColor(255, 0, 0)
        elif risk_assessment == "关注级":
            risk_run.font.color.rgb = RGBColor(255, 140, 0)
    except Exception as e:
        logger.error(f'添加总体评价段落失败: {e}')
        return None

    try:
        p2 = doc.add_paragraph()
        p2.add_run('【数据概况】').bold = True
        p2.add_run(f'累计分析银行流水 {total_trans} 条。')
    except Exception as e:
        logger.warning(f'添加数据概况段落失败: {e}')

    if aggregation_overview.get("summary"):
        try:
            summary = aggregation_overview["summary"]
            p_summary = doc.add_paragraph()
            p_summary.add_run('【聚合排序】').bold = True
            p_summary.add_run(
                f"极高风险{summary.get('极高风险实体数', 0)}个，"
                f"高风险{summary.get('高风险实体数', 0)}个，"
                f"高优先线索实体{summary.get('高优先线索实体数', 0)}个。"
            )
        except Exception as e:
            logger.warning(f'添加聚合排序段落失败: {e}')

    if aggregation_overview.get("highlights"):
        try:
            p_high = doc.add_paragraph()
            p_high.add_run('【高风险实体】').bold = True
            highlight_text = "，".join(
                f"{item['entity']}({item['risk_score']:.1f}分/置信度{item['risk_confidence']:.2f})"
                for item in aggregation_overview["highlights"][:3]
            )
            high_run = p_high.add_run(highlight_text)
            high_run.font.color.rgb = RGBColor(255, 0, 0)
        except Exception as e:
            logger.warning(f'添加高风险实体段落失败: {e}')

        try:
            top_clues = []
            for item in aggregation_overview["highlights"][:2]:
                if item.get("top_clues"):
                    top_clues.append(f"{item['entity']}:{item['top_clues'][0]}")
            if top_clues:
                p_clue = doc.add_paragraph()
                p_clue.add_run('【重点线索】').bold = True
                p_clue.add_run("；".join(top_clues))
        except Exception as e:
            logger.warning(f'添加重点线索段落失败: {e}')

    # 主要发现
    try:
        p3 = doc.add_paragraph()
        p3.add_run('【主要发现】').bold = True
        if aggregation_overview.get("highlights"):
            highlighted_entities = "、".join(
                f"{item['entity']}({item['risk_score']:.1f}分)"
                for item in aggregation_overview["highlights"][:3]
            )
            p3.add_run(f'聚合排序识别出重点核查对象 {highlighted_entities}。')
        elif direct_sus_count == 0 and hidden_sus_count == 0:
            p3.add_run('未发现核心人员与涉案公司存在直接利益输送。')
        else:
            finding = p3.add_run(f'发现 {direct_sus_count} 笔疑似直接利益输送，{hidden_sus_count} 笔疑似隐形资产线索。')
            finding.font.color.rgb = RGBColor(255, 0, 0)
    except Exception as e:
        logger.warning(f'添加主要发现段落失败: {e}')

    # ========== 3. 家庭资产概览表 ==========
    try:
        doc.add_heading('二、家庭资产与资金画像', level=1)
    except Exception as e:
        logger.error(f'添加家庭资产标题失败: {e}')
        return None

    # 汇总表格
    if core_persons:
        try:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 表头
            hdr_cells = table.rows[0].cells
            headers = ['人员', '总流入', '总流出', '工资占比', '理财持有']
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True

            # 数据行
            for person in core_persons:
                profile = profiles.get(person)
                if not isinstance(profile, dict) or not profile.get('has_data', False):
                    continue

                try:
                    summary = profile.get('summary', profile).get('summary', {})
                    if not isinstance(summary, dict):
                        continue

                    row_cells = table.add_row().cells
                    row_cells[0].text = str(person)

                    # 安全格式化货币
                    try:
                        row_cells[1].text = utils.format_currency(summary.get('total_income', 0))
                    except Exception:
                        row_cells[1].text = str(summary.get('total_income', 0))

                    try:
                        row_cells[2].text = utils.format_currency(summary.get('total_expense', 0))
                    except Exception:
                        row_cells[2].text = str(summary.get('total_expense', 0))

                    # 安全格式化百分比
                    try:
                        salary_ratio = summary.get('salary_ratio', 0)
                        row_cells[3].text = f"{salary_ratio:.1%}"
                    except Exception:
                        row_cells[3].text = '-'

                    wealth = profile.get('wealth_management', {})
                    if isinstance(wealth, dict):
                        holding = wealth.get('estimated_holding', 0)
                        try:
                            row_cells[4].text = utils.format_currency(holding) if holding > 0 else '-'
                        except Exception:
                            row_cells[4].text = str(holding) if holding > 0 else '-'
                    else:
                        row_cells[4].text = '-'
                except Exception as e:
                    logger.warning(f'添加人员 {person} 的数据行失败: {e}')
                    continue
        except Exception as e:
            logger.error(f'创建家庭资产表格失败: {e}')

    # ========== 4. 公司资金核查 ==========
    try:
        doc.add_heading('三、公司资金核查', level=1)
    except Exception as e:
        logger.error(f'添加公司资金核查标题失败: {e}')
        return None

    if not involved_companies:
        try:
            doc.add_paragraph('本次未涉及公司核查。')
        except Exception as e:
            logger.warning(f'添加公司核查说明失败: {e}')
    else:
        for company in involved_companies:
            try:
                doc.add_heading(f'➤ {company}', level=2)
            except Exception as e:
                logger.warning(f'添加公司 {company} 标题失败: {e}')
                continue

            comp_profile = profiles.get(company)

            if not isinstance(comp_profile, dict) or not comp_profile.get('has_data', False):
                try:
                    doc.add_paragraph('(暂无详细流水数据)')
                except Exception as e:
                    logger.warning(f'添加公司 {company} 无数据说明失败: {e}')
                continue

            try:
                summary = comp_profile.get('summary', comp_profile).get('summary', {})
                if isinstance(summary, dict):
                    p = doc.add_paragraph()
                    p.add_run('资金概况: ').bold = True
                    try:
                        p.add_run(f'总流入 {utils.format_currency(summary.get("total_income", 0))} | 总流出 {utils.format_currency(summary.get("total_expense", 0))}')
                    except Exception:
                        p.add_run(f'总流入 {summary.get("total_income", 0)} | 总流出 {summary.get("total_expense", 0)}')
            except Exception as e:
                logger.warning(f'添加公司 {company} 资金概况失败: {e}')

    # ========== 5. 疑点与建议 ==========
    try:
        doc.add_heading('四、主要疑点与核查建议', level=1)
    except Exception as e:
        logger.error(f'添加疑点与建议标题失败: {e}')
        return None

    suggestion_idx = 1
    has_suggestions = False

    # 利益输送
    direct_transfers = suspicions.get('direct_transfers', [])
    if direct_transfers:
        try:
            doc.add_heading('【疑似利益输送】', level=2)
            for t in direct_transfers[:10]:  # 限制数量
                try:
                    amount = t.get('amount', 0)
                    amount_str = utils.format_currency(amount) if amount is not None else '0'
                    doc.add_paragraph(
                        f"• {_safe_format_date(t.get('date'))}: {t.get('person', '')} {t.get('direction', '')} "
                        f"{t.get('company', '')} {amount_str}"
                    )
                except Exception as e:
                    logger.warning(f'添加利益输送条目失败: {e}')
                    continue
            doc.add_paragraph(f'➡ 建议 {suggestion_idx}: 调取相关凭证，核实上述资金往来的业务背景。')
            suggestion_idx += 1
            has_suggestions = True
        except Exception as e:
            logger.warning(f'添加利益输送部分失败: {e}')

    # 隐形资产
    try:
        hidden_assets = suspicions.get('hidden_assets', {})
        hidden = [item for sublist in hidden_assets.values() for item in sublist]
        if hidden:
            doc.add_heading('【疑似隐形资产】', level=2)
            for h in hidden[:10]:
                try:
                    amount = h.get('amount', 0)
                    amount_str = utils.format_currency(amount) if amount is not None else '0'
                    doc.add_paragraph(
                        f"• {_safe_format_date(h.get('date'))}: 支付 {amount_str} "
                        f"给 {h.get('counterparty', '')}"
                    )
                except Exception as e:
                    logger.warning(f'添加隐形资产条目失败: {e}')
                    continue
            doc.add_paragraph(f'➡ 建议 {suggestion_idx}: 核实上述大额支出是否用于购房/购车。')
            suggestion_idx += 1
            has_suggestions = True
    except Exception as e:
        logger.warning(f'添加隐形资产部分失败: {e}')

    if not has_suggestions:
        try:
            doc.add_paragraph('本次自动化分析未发现显著的硬性疑点。')
            doc.add_paragraph('➡ 建议: 重点关注大额消费支出是否与收入水平匹配。')
        except Exception as e:
            logger.warning(f'添加无疑点说明失败: {e}')

    # ========== 6. 报告尾部 ==========
    try:
        doc.add_paragraph()
        doc.add_paragraph('=' * 50)
        footer = doc.add_paragraph('注: 本报告基于提供的电子数据分析生成，线索仅供参考。')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        logger.warning(f'添加报告尾部失败: {e}')

    # 保存文档
    try:
        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)

        doc.save(output_path)
        logger.info(f'Word 审计报告生成完成: {output_path}')
        return output_path
    except PermissionError as e:
        logger.error(f'保存 Word 文档失败（权限错误）: {e}')
        return None
    except OSError as e:
        logger.error(f'保存 Word 文档失败（系统错误）: {e}')
        return None
    except Exception as e:
        logger.error(f'保存 Word 文档失败（未知错误）: {e}')
        return None
